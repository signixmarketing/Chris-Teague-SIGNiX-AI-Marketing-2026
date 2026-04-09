#!/usr/bin/env python3
"""Build SIGNiX_PaidMedia_4.6.26.docx — Rev 3, April 6 2026."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Brand colours ─────────────────────────────────────────────────────────────
GREEN = RGBColor(0x6d, 0xa3, 0x4a)
INK   = RGBColor(0x2e, 0x34, 0x40)
BODY  = RGBColor(0x54, 0x54, 0x54)
WHITE = RGBColor(0xff, 0xff, 0xff)
FONT  = "Calibri"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def style_run(run, color, bold=False, size=11, italic=False):
    run.font.name      = FONT
    run.font.color.rgb = color
    run.font.bold      = bold
    run.font.size      = Pt(size)
    run.font.italic    = italic

def _apply_inline(p, text, base_bold=False, base_italic=False, base_color=None):
    color = base_color or BODY
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            style_run(run, INK, bold=True, size=11)
        else:
            run = p.add_run(part)
            style_run(run, color, bold=base_bold, size=11, italic=base_italic)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=24)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=16)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "6da34a")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, GREEN, bold=True, size=12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    _apply_inline(p, text, base_italic=italic)
    p.paragraph_format.space_after = Pt(6)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    _apply_inline(p, text)

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "2e3440")
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        style_run(run, WHITE, bold=True, size=10)
    for r_i, row in enumerate(rows):
        bg = "f8fafb" if r_i % 2 == 0 else "ffffff"
        for c_i, val in enumerate(row):
            c = tbl.rows[r_i + 1].cells[c_i]
            set_cell_bg(c, bg)
            c.paragraphs[0].clear()
            _apply_inline(c.paragraphs[0], val)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "6da34a")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    style_run(run, BODY, italic=True, size=11)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "d8dee9")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── BUILD ─────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(1.0)
sec.right_margin  = Inches(1.0)
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)

doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(11)

# ── Title ─────────────────────────────────────────────────────────────────────
add_h1(doc, "SIGNiX Paid Media Plan — 2026")

meta_rows = [
    ("Prepared by", "Chris, Head of Growth and Marketing"),
    ("Date",        "April 6, 2026"),
    ("Status",      "CEO-approved — active. Updated April 6, 2026 (Rev 3)."),
    ("Budget",      "$3,000/month — Google Search Ads only. LinkedIn approved as Phase 2 when 60-day Google performance is proven."),
    ("Ad scheduling", "8:00 am – 5:00 pm EST, Monday–Friday. Ads paused outside these hours so leads get same-day follow-up."),
    ("Objective",   "Shift paid media from individual notary leads to high-value B2B pipeline. Three target sectors: financial services / healthcare, wealth management, and legal / debt collection — all places where a disputed signature is a legal or financial event."),
]
mt = doc.add_table(rows=len(meta_rows), cols=2)
mt.style = "Table Grid"
for i, (k, v) in enumerate(meta_rows):
    set_cell_bg(mt.rows[i].cells[0], "2e3440")
    mt.rows[i].cells[0].width = Inches(1.5)
    run_k = mt.rows[i].cells[0].paragraphs[0].add_run(k)
    style_run(run_k, WHITE, bold=True, size=10)
    set_cell_bg(mt.rows[i].cells[1], "f8fafb")
    run_v = mt.rows[i].cells[1].paragraphs[0].add_run(v)
    style_run(run_v, BODY, size=10)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_hr(doc)

# ── Section 1 ─────────────────────────────────────────────────────────────────
add_h2(doc, "1. The Case for Reallocation")
add_h3(doc, "Current state")
add_body(doc, "SIGNiX's paid media spend was concentrated on Google Ads targeting individual notary and RON queries. These campaigns attract individual notary professionals — low ACV, high-touch customers who require a full sales close plus dedicated CSM onboarding.")
add_body(doc, "**The economics are broken at the unit level.** When the salary cost of a sales rep and CSM is divided by the time spent per notary lead, the contribution margin per notary customer is likely near zero or negative.")
add_bullet(doc, "Individual notary deals: low ACV, one-time onboarding cost, low recurring transaction volume")
add_bullet(doc, "Staff time per deal: too high relative to deal value")
add_bullet(doc, "Scalability: none — each deal requires the same manual motion regardless of volume")

add_h3(doc, "Proposed state")
add_body(doc, "Redirect the same $3,000/month toward buyers in sectors where a disputed or fraudulent signature is a legal and financial event:")
add_bullet(doc, "**Financial services / Healthcare:** Banks, credit unions, broker-dealers, RIAs, and health systems under FINRA/SEC/HIPAA requirements. A failed audit costs more than years of SIGNiX.")
add_bullet(doc, "**Wealth management:** RIAs and financial advisors handling beneficiary changes, IRA distributions, and account events where a disputed signature triggers legal and financial consequences.")
add_bullet(doc, "**Legal / Debt collection:** Law firms and collection agencies under FDCPA and state consent requirements. Proving who signed is the difference between a valid agreement and a disputed one.")
add_body(doc, "**Note:** Real estate / mortgage was removed as a target sector. New mortgage customer acquisition faces too many barriers — regulatory complexity, long sales cycles, and dominant vendor relationships — to be a productive paid media target at this budget level.", italic=True)
add_hr(doc)

# ── Section 2 ─────────────────────────────────────────────────────────────────
add_h2(doc, "2. Budget Allocation")
add_table(doc,
    headers=["Channel", "Monthly Allocation", "Status", "Rationale"],
    rows=[
        ["Google Search Ads",        "$3,000", "Active",              "Capture buyers actively searching in financial services, healthcare, wealth management, and legal/debt collection"],
        ["LinkedIn Sponsored Content", "TBD",  "Phase 2 — on hold",  "Launch after 60-day Google review. Requires landing pages and proven CPL before adding a second channel."],
        ["**Total**",               "**$3,000**", "",                 "No increase from current spend — reallocation only"],
    ],
    col_widths=[1.8, 1.3, 1.3, 2.1]
)
add_body(doc, "Ad scheduling: All campaigns run 8:00 am – 5:00 pm EST, Monday–Friday. Ads are paused outside these hours to avoid wasted spend on leads no one can act on immediately.", italic=True)
add_hr(doc)

# ── Section 3 ─────────────────────────────────────────────────────────────────
add_h2(doc, "3. Google Search Ads Strategy ($3,000/month)")
add_h3(doc, "Strategic principle")
add_body(doc, "Avoid head-to-head competition on generic e-sign terms where DocuSign and Adobe dominate with multi-million-dollar budgets. Compete on specificity: regulated industry, compliance urgency, authentication, and workflow-level search intent.")
add_body(doc, "**The ID Verify filter:** Before adding any keyword, ask — if this document gets disputed, does the signer's identity matter? If yes, SIGNiX belongs there and ID Verify is the differentiator.")

add_h3(doc, "Bucket 1 — Compliance / Regulatory (~$800/month) — PRIORITY")
add_body(doc, "Target: CCOs, compliance managers, and VP Operations at broker-dealers, RIAs, banks, and credit unions who are actively searching because of a FINRA exam, SEC audit, or internal compliance finding.")
add_table(doc,
    headers=["Keyword", "Intent level"],
    rows=[
        ["`SEC Rule 17a-4 compliant e-signature`",          "High — very specific, low competition"],
        ["`FINRA compliant digital signature`",              "High — exam-triggered intent"],
        ["`FINRA electronic signature audit trail`",         "High"],
        ["`electronic recordkeeping compliance broker dealer`", "High"],
        ["`non-repudiation e-signature financial services`", "High — advanced buyer signal"],
        ["`e-signature FINRA Rule 4511`",                    "High intent, minimal competition"],
        ["`digital signature SEC compliance`",               "High"],
    ],
    col_widths=[3.5, 3.0]
)
add_body(doc, "Match type: Phrase and exact only.")

add_h3(doc, "Bucket 2 — Authentication / Fraud (~$600/month)")
add_body(doc, "Target: FI ops and compliance leaders thinking about fraud risk. Connects directly to Mike's email campaign and SIGNiX Fraud Alert.")
add_table(doc,
    headers=["Keyword", "Intent level", "Notes"],
    rows=[
        ["`identity verification electronic signature`",     "High — core SIGNiX differentiator", "Revised phrasing — better search match than tested variant"],
        ["`e-signature with identity verification`",         "High — core SIGNiX differentiator", ""],
        ["`verified electronic signature`",                  "High",                              "Broad but intent-rich"],
        ["`electronic signature signer authentication`",     "High",                              ""],
        ["`electronic signature fraud prevention banking`",  "High",                              ""],
        ["`two-factor authentication electronic signature`", "High",                              ""],
        ["`biometric e-signature financial services`",       "Medium — emerging search behavior", ""],
        ["`signer authentication platform`",                 "High",                              "Workflow-level intent"],
    ],
    col_widths=[2.7, 1.7, 2.1]
)

add_h3(doc, "Bucket 3 — Wealth Management / Financial Advisor Workflows (~$500/month)")
add_body(doc, "Target: RIAs, broker-dealers, and financial advisors executing documents where a disputed identity has direct financial or legal consequence. Focus on workflow-specific search terms, not product feature terms.")
add_table(doc,
    headers=["Keyword", "Intent level", "Notes"],
    rows=[
        ["`beneficiary change form`",                        "High — **390 avg monthly searches, +50% YoY**", "Validated top performer. Cheap bids ($1–$4). Anchor term."],
        ["`beneficiary change form e-signature`",            "High",                                          "Companion term to above"],
        ["`power of attorney e-signature`",                  "High",                                          "Irrevocable authority document — ID Verify critical"],
        ["`POA document signing platform`",                  "High",                                          "Workflow-level search intent"],
        ["`investment account opening signature`",           "High",                                          "KYC and fraud risk — ID Verify moment"],
        ["`401k distribution electronic signature`",         "High",                                          "Tax consequence — identity verification warranted"],
        ["`wealth management document signing`",             "Medium — good volume",                          ""],
        ["`financial advisor e-signature compliance`",       "High",                                          ""],
        ["`e-signature for investment advisors`",            "High",                                          ""],
        ["`SEC registered investment adviser e-sign`",       "High intent, low competition",                  ""],
    ],
    col_widths=[2.5, 1.9, 2.1]
)

add_h3(doc, "Bucket 4 — Industry / Use-Case (~$600/month)")
add_body(doc, "Target: FI ops, IT evaluators, healthcare administrators, and legal/collections professionals with active workflow problems. Three sub-segments: financial services, healthcare, and legal/debt collection.")

add_body(doc, "**Financial services:**")
add_table(doc,
    headers=["Keyword", "Intent level", "Notes"],
    rows=[
        ["`e-signature for credit unions`",              "High — lower competition",     ""],
        ["`loan document electronic signature`",         "High — very specific",         ""],
        ["`electronic signature NCUA compliant`",        "High intent, very low comp.",  ""],
        ["`new account opening electronic signature`",   "High",                         "KYC and fraud risk — ID Verify moment"],
        ["`wire transfer authorization form`",           "High",                         "Highest fraud-risk signing moment"],
        ["`wire transfer form electronic signature`",    "High",                         ""],
        ["`ACH authorization e-signature`",              "High",                         "Payment fraud risk"],
    ],
    col_widths=[2.5, 1.7, 2.3]
)

add_body(doc, "**Healthcare (HIPAA):**")
add_table(doc,
    headers=["Keyword", "Intent level", "Notes"],
    rows=[
        ["`HIPAA compliant e-signature`",                "High — **210/month, +191% in 3 months**", "Validated top performer. Already in account. Anchor term."],
        ["`HIPAA e-signature`",                          "High",                                     "Shorter variation — likely higher volume"],
        ["`HIPAA electronic signature`",                 "High",                                     "Alternate phrasing — test alongside above"],
        ["`patient consent electronic signature`",       "High — HIPAA workflow",                    ""],
        ["`telehealth consent form signature`",          "High — growing segment",                   ""],
        ["`medical records release e-signature`",        "High",                                     ""],
    ],
    col_widths=[2.5, 1.9, 2.1]
)

add_body(doc, "**Legal / Debt collection:**")
add_table(doc,
    headers=["Keyword", "Intent level", "Notes"],
    rows=[
        ["`FDCPA compliant e-signature`",                        "High — low competition",   "FDCPA is SIGNiX's FINRA equivalent for this sector"],
        ["`debt settlement agreement electronic signature`",     "High",                     "Consent documentation failure = legal exposure"],
        ["`debt collection consent form`",                       "High",                     ""],
        ["`legal document e-signature platform`",                "Medium — broader",         ""],
        ["`collection agency e-signature`",                      "High",                     ""],
        ["`attorney e-signature compliance`",                    "High — low competition",   ""],
    ],
    col_widths=[2.5, 1.5, 2.5]
)

add_h3(doc, "Bucket 5 — Competitor Displacement (~$400/month)")
add_body(doc, "Target: FI evaluators actively comparing solutions — mid-funnel buyers who are already looking.")
add_table(doc,
    headers=["Keyword", "Intent level", "Notes"],
    rows=[
        ["`docusign alternative`",                        "High — moderate volume",   "Broader term; test alongside sector-specific variants"],
        ["`docusign alternative for banks`",              "High",                     ""],
        ["`docusign alternative financial services`",     "High",                     ""],
        ["`docusign alternative credit unions`",          "High — lower CPC",         ""],
        ["`adobe sign alternative compliance`",           "High",                     ""],
        ["`docusign alternative compliance`",             "High",                     ""],
    ],
    col_widths=[2.5, 1.5, 2.5]
)
add_body(doc, "Match type: Phrase and exact. Broad match excluded — attracts unqualified traffic on DocuSign brand terms.")

add_h3(doc, "Bucket 6 — RON Institutional only (~$100/month)")
add_body(doc, "Drop all individual notary terms and any terms confirmed at zero volume. Budget is minimal — maintenance only, not growth.")
add_table(doc,
    headers=["Keyword", "Intent level", "Notes"],
    rows=[
        ["`RON platform compliance`",            "High", "Retained — institutional buyer intent"],
        ["`remote online notary for banks`",     "High — low competition", "Replacement for removed zero-volume term"],
    ],
    col_widths=[2.5, 1.5, 2.5]
)

add_h3(doc, "Ad message principles")
add_bullet(doc, "**Lead with fear. Close with partnership.** The buyer's fear opens the door. The human relationship closes the deal.")
add_bullet(doc, "Fear-first openers: 'One disputed signature you can't prove costs more than years of SIGNiX.' / 'FINRA is flagging electronic records right now. Can you prove who signed?' / 'AI fraud is targeting financial accounts. A click is not a signature.'")
add_bullet(doc, "Partnership close: 'SIGNiX assigns you a CSM who knows your workflows before a problem happens.'")
add_bullet(doc, "Each bucket lands on a dedicated landing page matched to the search intent — not the generic homepage.")
add_bullet(doc, "Never lead with PKI, features, or product specs — lead with the consequence of not having it.")
add_hr(doc)

# ── Section 3a ────────────────────────────────────────────────────────────────
add_h2(doc, "3a. Keywords Tested and Removed")
add_body(doc, "The following keywords were submitted to Google Keyword Planner for 12-month volume analysis (March 2025 – February 2026). All returned zero or near-zero search volume and have been removed from active campaigns.")
add_table(doc,
    headers=["Keyword", "Former Bucket", "Avg Monthly Searches", "YoY Change", "Reason Removed"],
    rows=[
        ["`e signature for banks`",                          "Industry / Use-Case",    "20",  "–75%", "Dying term. Volume declined sharply year over year."],
        ["`digital signature mortgage docs`",                "Industry / Use-Case",    "0",   "N/A",  "Zero volume. Mortgage sector also removed as a target."],
        ["`remote online notary platform financial inst.`",  "RON Institutional",      "0",   "N/A",  "Zero volume. Too niche as a search phrase."],
        ["`e signature rollover forms`",                     "Wealth Management",      "0",   "N/A",  "Zero volume. Buyers search the document type, not the signing action."],
        ["`ira distribution e sign`",                        "Wealth Management",      "0",   "N/A",  "Zero volume. Industry shorthand — not how buyers search."],
        ["`kba electronic signature`",                       "Authentication / Fraud", "0",   "N/A",  "Zero volume. 'KBA' is internal industry jargon. Buyers don't know the acronym."],
        ["`e signature identity verification`",              "Authentication / Fraud", "0",   "N/A",  "Zero volume. Wrong word order for how buyers search. Replaced with `identity verification electronic signature`."],
        ["`ai fraud e signature`",                           "Authentication / Fraud", "0",   "N/A",  "Zero volume. Search behavior for this topic has not yet developed. Re-test Q3 2026."],
    ],
    col_widths=[1.7, 1.3, 0.9, 0.7, 1.9]
)
add_body(doc, "**What the data confirmed:** The two keywords with validated volume were `beneficiary change form` (390/month, +50% YoY) and `HIPAA compliant e-signature` (210/month, +191% in 3 months). Both are retained and treated as anchor terms for their respective buckets.")
add_hr(doc)

# ── Section 4 ─────────────────────────────────────────────────────────────────
add_h2(doc, "4. LinkedIn Strategy — Phase 2 (On Hold)")
add_body(doc, "LinkedIn is approved in principle but on hold until Google performance is proven. Target launch: 60-day Google review (June 2026).")
add_h3(doc, "Why LinkedIn matters when we're ready")
add_body(doc, "FI decision-makers and ISV product leaders do not frequently Google for e-sign solutions unprompted. They respond to peer content, authority positioning, and problem-aware messaging in their professional feed. LinkedIn reaches cold but qualified buyers before a search event occurs.")

add_h3(doc, "Planned Audience 1 — FI operations, compliance, and wealth management")
add_table(doc,
    headers=["Targeting dimension", "Values"],
    rows=[
        ["Job titles",    "VP Operations, Chief Compliance Officer, Director of Lending, Financial Advisor, Wealth Manager, Director of Digital Banking"],
        ["Industries",    "Banking, Financial Services, Credit Unions, Investment Management"],
        ["Company size",  "50–5,000 employees"],
        ["Geography",     "United States"],
        ["Message angle", "Fear-first: 'One disputed transaction you can't prove costs more than years of SIGNiX.' Partnership close: 'We assign you a CSM who knows your workflow.'"],
    ],
    col_widths=[1.8, 4.7]
)

add_h3(doc, "Planned Audience 2 — ISV and Flex API targets")
add_table(doc,
    headers=["Targeting dimension", "Values"],
    rows=[
        ["Job titles",    "Product Manager, VP of Product, CTO, VP Engineering, Head of Product"],
        ["Industries",    "Financial Services Software, Fintech, Information Technology"],
        ["Company size",  "10–500 employees"],
        ["Geography",     "United States"],
        ["Message angle", "'Your customers live in regulated industries. Compliance is in the SIGNiX Flex API — not your problem to build.'"],
    ],
    col_widths=[1.8, 4.7]
)
add_hr(doc)

# ── Section 5 ─────────────────────────────────────────────────────────────────
add_h2(doc, "5. Landing Pages (prerequisite for launch)")
add_body(doc, "Each keyword bucket should land on a page that matches the search intent. Sending all paid traffic to the homepage is the single biggest waste of paid media budget.")
add_table(doc,
    headers=["Traffic source", "Recommended landing page focus"],
    rows=[
        ["Compliance / Regulatory bucket",    "'FINRA exam coming? Here's how SIGNiX proves who signed.'"],
        ["Authentication / Fraud bucket",     "'Anyone can click a button. SIGNiX proves it was your customer.'"],
        ["Wealth Management bucket",          "'Built for the forms advisors use every day — with an audit trail.'"],
        ["Financial services (industry)",     "'Electronic signature for banks and credit unions — compliance-ready.'"],
        ["Healthcare (HIPAA)",               "'HIPAA requires you to prove who consented. SIGNiX does that.'"],
        ["Legal / Debt collection",           "'FDCPA requires consent documentation. SIGNiX gives you proof.'"],
        ["Competitor displacement",           "'How SIGNiX compares to DocuSign for regulated industries.'"],
    ],
    col_widths=[2.2, 4.3]
)
add_body(doc, "Build order: Start with two pages — compliance/financial services and healthcare/HIPAA. Add legal/debt collection and competitor displacement pages in weeks 3–4.", italic=True)
add_hr(doc)

# ── Section 6 ─────────────────────────────────────────────────────────────────
add_h2(doc, "6. Leading and Lagging Indicators")
add_h3(doc, "Leading indicators — check weekly")
add_table(doc,
    headers=["Indicator", "What it tells you", "Target"],
    rows=[
        ["Impression share by keyword bucket",          "Are we visible on priority terms?",           ">40% on priority buckets"],
        ["CTR by bucket",                               "Are the right people clicking?",              "3–6% for B2B search"],
        ["Average CPC trends",                          "Is spend efficient or inflating?",            "Monitor week-over-week"],
        ["Landing page bounce rate",                    "Is traffic finding what they expected?",      "<60%"],
        ["Form fill / demo request rate",               "Are visitors converting?",                    "2–5% of paid clicks"],
        ["Lead qualification rate",                     "Are leads matching ICP?",                     ">40% qualified"],
        ["Time from lead to first sales touchpoint",    "Is the handoff working?",                     "<48 hours"],
    ],
    col_widths=[2.2, 2.3, 2.0]
)

add_h3(doc, "Lagging indicators — check monthly and quarterly")
add_table(doc,
    headers=["Indicator", "What it tells you", "Notes"],
    rows=[
        ["Cost per qualified lead (CPL) by bucket",         "Which bucket delivers the most efficient leads?",   "Split by bucket"],
        ["MQLs per month",                                   "Is pipeline building?",                             "Define MQL criteria before launch"],
        ["Sales-accepted leads (SALs)",                      "Is the sales team seeing quality?",                 "Track rep feedback per lead source"],
        ["Pipeline sourced ($) from paid channels",          "What's the revenue potential?",                     "Attribute in HubSpot"],
        ["Deals closed from paid sources",                   "What's the actual return?",                         "90–180 day lag expected"],
        ["Average deal ACV by lead source",                  "Are paid leads higher value?",                      "Compare vs. notary lead ACV"],
        ["Transaction volume from new paid-sourced customers", "Does it generate recurring revenue?",             "The ultimate SIGNiX metric"],
    ],
    col_widths=[2.2, 2.3, 2.0]
)
add_hr(doc)

# ── Section 7 ─────────────────────────────────────────────────────────────────
add_h2(doc, "7. 90-Day Launch Plan")
add_table(doc,
    headers=["Week", "Action", "Owner"],
    rows=[
        ["1–2",  "Pause all individual notary keyword spending; keep RON institutional terms at $100/month",               "Chris"],
        ["1–2",  "Build compliance/FI landing page and healthcare/HIPAA landing page (two pages minimum before launch)",   "Chris + web"],
        ["1–2",  "Set up HubSpot lead source tracking for paid search; confirm daily lead monitoring is in place",         "Chris"],
        ["3–4",  "Launch Google Buckets 1–4 (compliance, authentication, wealth management, industry/use-case) at 80%",   "Chris"],
        ["3–4",  "Confirm ad scheduling is set to 8am–5pm EST in Google Ads campaign settings",                           "Chris"],
        ["5–6",  "Review CTR and lead quality by bucket; pause underperforming ad variants; add legal/debt collection page","Chris"],
        ["6–8",  "Launch Bucket 5 (competitor displacement); begin planning LinkedIn Phase 2 based on Google data",        "Chris"],
        ["8–12", "Review CPL by bucket; consolidate spend into top two or three buckets",                                   "Chris"],
        ["12",   "First 90-day report: CPL, pipeline sourced, SALs, ACV vs. notary benchmark; present LinkedIn Phase 2",  "Chris → CEO"],
    ],
    col_widths=[0.6, 5.2, 0.8]
)
add_hr(doc)

# ── Section 8 ─────────────────────────────────────────────────────────────────
add_h2(doc, "8. Rationale Summary")
add_blockquote(doc,
    "We were spending money on leads that cost us more to close and onboard than they generate. We stopped that. "
    "With $3,000/month — Google only, running 8am to 5pm EST — we are going after buyers in three sectors where a "
    "disputed or fraudulent signature is a legal and financial event: financial services and healthcare, wealth "
    "management, and legal and debt collection. These buyers are searching right now. FINRA is flagging electronic "
    "records in active exams. HIPAA searches are up 191% in three months. FDCPA consent failures are a growing "
    "exposure for collections teams. Our message is simple: a click is not a signature. SIGNiX proves who signed. "
    "And unlike our competitors, we assign every customer a CSM who knows their workflow. We removed eight keywords "
    "with zero search volume in April 2026 and replaced them with workflow-specific terms where identity verification "
    "changes the outcome. We will know if this is working within 60 days."
)
add_hr(doc)
add_body(doc, "Last updated: April 6, 2026 — Rev 3. Updated to reflect keyword volume analysis (March 2025–February 2026): eight zero-volume keywords removed and documented in Section 3a; mortgage sector replaced by legal/debt collection; expanded keyword list across all buckets with workflow-specific, ID Verify-aligned terms.", italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/chris/Desktop/AI Summit 2026 Q2/proj-template-and-lease-SIGNiX-app/PROJECT-DOCS/SIGNiX_PaidMedia_4.6.26.docx"
doc.save(out)
print(f"Saved: {out}")
