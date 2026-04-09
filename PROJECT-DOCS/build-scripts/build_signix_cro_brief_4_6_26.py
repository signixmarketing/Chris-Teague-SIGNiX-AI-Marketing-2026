#!/usr/bin/env python3
"""Build SIGNiX_CRO_PaidMedia_Brief_4.6.26.docx — CRO one-page brief."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Brand colours ─────────────────────────────────────────────────────────────
GREEN = RGBColor(0x6d, 0xa3, 0x4a)
INK   = RGBColor(0x2e, 0x34, 0x40)
BODY  = RGBColor(0x54, 0x54, 0x54)
MUTED = RGBColor(0x6b, 0x72, 0x80)
WHITE = RGBColor(0xff, 0xff, 0xff)
FONT  = "Calibri"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def style_run(run, color, bold=False, size=11, italic=False):
    run.font.name      = FONT
    run.font.color.rgb = color
    run.font.bold      = bold
    run.font.size      = Pt(size)
    run.font.italic    = italic

def _apply_inline(p, text, base_bold=False, base_italic=False, base_color=None):
    color = base_color or BODY
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            style_run(run, INK, bold=True, size=10)
        else:
            run = p.add_run(part)
            style_run(run, color, bold=base_bold, size=10, italic=base_italic)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=13)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "6da34a")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, GREEN, bold=True, size=10)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

def add_body(doc, text, italic=False, size=10):
    p = doc.add_paragraph()
    _apply_inline(p, text, base_italic=italic, base_color=BODY)
    p.paragraph_format.space_after = Pt(4)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    _apply_inline(p, text)

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "2e3440")
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        style_run(run, WHITE, bold=True, size=9)
    for r_i, row in enumerate(rows):
        bg = "f8fafb" if r_i % 2 == 0 else "ffffff"
        for c_i, val in enumerate(row):
            c = tbl.rows[r_i + 1].cells[c_i]
            set_cell_bg(c, bg)
            c.paragraphs[0].clear()
            _apply_inline(c.paragraphs[0], val)
            c.paragraphs[0].paragraph_format.space_after = Pt(1)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return tbl

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "6da34a")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    style_run(run, BODY, italic=True, size=10)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "d8dee9")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)


# ── BUILD ─────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(0.9)
sec.right_margin  = Inches(0.9)
sec.top_margin    = Inches(0.8)
sec.bottom_margin = Inches(0.8)

doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(10)

# ── Header ────────────────────────────────────────────────────────────────────
add_h1(doc, "SIGNiX Paid Media — What This Means for the Sales Team")

meta_rows = [
    ("Prepared for", "Chief Revenue Officer — SIGNiX"),
    ("Prepared by",  "Chris, Head of Growth and Marketing"),
    ("Date",         "April 6, 2026"),
    ("Purpose",      "Set expectations, address concerns about lead volume, and show how sales wins with this approach."),
]
mt = doc.add_table(rows=len(meta_rows), cols=2)
mt.style = "Table Grid"
for i, (k, v) in enumerate(meta_rows):
    set_cell_bg(mt.rows[i].cells[0], "2e3440")
    mt.rows[i].cells[0].width = Inches(1.3)
    run_k = mt.rows[i].cells[0].paragraphs[0].add_run(k)
    style_run(run_k, WHITE, bold=True, size=9)
    set_cell_bg(mt.rows[i].cells[1], "f8fafb")
    run_v = mt.rows[i].cells[1].paragraphs[0].add_run(v)
    style_run(run_v, BODY, size=9)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_hr(doc)

# ── Section 1: The Short Version ──────────────────────────────────────────────
add_h2(doc, "The Short Version")
add_body(doc,
    "We changed the keyword strategy in April. The goal is fewer, higher-quality leads. "
    "Not more volume. This brief explains what that means for the sales team and what you "
    "can expect in the next 60 days.")
add_body(doc,
    "The concern about lead volume is fair. Here is why the numbers still work.")
add_hr(doc)

# ── Section 2: Why We Made This Change ───────────────────────────────────────
add_h2(doc, "Why We Changed the Approach")
add_body(doc,
    "The previous keyword strategy used broad e-signature terms. Those terms attract "
    "small businesses and price shoppers. Our average closed deal is **$4,100**. "
    "That buyer does not search 'free DocuSign alternative.'")
add_body(doc,
    "We tested 75+ keywords in Google Keyword Planner. Here is what we found:")
add_table(doc,
    headers=["Finding", "What We Did"],
    rows=[
        ["**15 keywords returned zero search volume**",
         "Removed. No budget allocated to dead terms."],
        ["4 high-volume competitor terms attract price shoppers",
         "Excluded. 'Free DocuSign alternative' gets 1,900 searches/month. Not our buyer."],
        ["**2 keywords validated with real buyer intent**",
         "beneficiary change form (390/mo) and HIPAA compliant e-signature (210/mo). Both in campaign."],
        ["docusign competitors gets 1,000 searches/month",
         "Active. Strong intent. Already running."],
        ["~30 new workflow-specific keywords identified",
         "In test queue for Q2. Volume data being collected now."],
    ],
    col_widths=[2.8, 3.8]
)
add_body(doc,
    "The new keywords target buyers with a specific problem: a compliance officer who "
    "needs HIPAA-compliant signing, a wealth manager who needs a beneficiary form signed "
    "with an audit trail, a collection agency that needs consent forms that hold up in court. "
    "Those buyers are worth chasing.")
add_hr(doc)

# ── Section 3: What the Team Can Expect ──────────────────────────────────────
add_h2(doc, "What the Sales Team Can Expect")
add_body(doc,
    "Lead volume will be lower than a broad campaign. That is by design. "
    "Here is the realistic 90-day picture:")
add_table(doc,
    headers=["Timeframe", "Projected Qualified Leads", "Expected Closed Deals", "Revenue (ACV)"],
    rows=[
        ["Month 1–2", "2–4 / month", "0–1",  "Ramp phase. Baseline being set."],
        ["Month 3",   "4–8 / month", "1–3",  "$4,100–$12,300"],
        ["Quarter total", "12–24",   "2–4",  "$8,200–$16,400"],
    ],
    col_widths=[1.3, 1.6, 1.6, 2.1]
)
add_body(doc,
    "Two to three closed deals puts this channel in positive ROI for the quarter. "
    "That is the benchmark. Transaction revenue from authenticated signings adds on top of ACV.")
add_blockquote(doc,
    "Sending the team 50 unqualified leads at a 2% close rate produces 1 deal. "
    "Sending them 10 qualified leads at a 20% close rate produces 2 deals, with "
    "less time wasted. We are building toward the second scenario.")
add_hr(doc)

# ── Section 4: What the Team Does While Leads Ramp ───────────────────────────
add_h2(doc, "What the Team Does While Leads Ramp")
add_body(doc,
    "Paid media leads will build over 60 days. The sales team does not need to wait.")
add_body(doc,
    "We built an **ABM Activity Scorecard** so the team can track outbound prospecting "
    "alongside inbound leads. Both count toward activity. Both show up in the numbers.")
add_bullet(doc, "Each rep identifies 5–10 target accounts per week from the ICP list.")
add_bullet(doc, "Touches (email, LinkedIn, call) are logged in the scorecard.")
add_bullet(doc, "Meeting status, content sent, and next steps are tracked in one view.")
add_bullet(doc, "Scorecard is reviewed weekly. Activity is visible to leadership.")
add_body(doc,
    "This is not extra work. It replaces waiting for the phone to ring with a structured "
    "outbound motion that the team controls. The scorecard is ready now.")
add_hr(doc)

# ── Section 5: The Ask ────────────────────────────────────────────────────────
add_h2(doc, "One Ask from Marketing")
add_body(doc,
    "**Respond to every paid media lead within 24 hours.**")
add_body(doc,
    "A qualified lead that goes cold is $4,100 left on the table. "
    "The keywords we are running attract buyers with a specific problem. "
    "Those buyers are also talking to one or two competitors. "
    "Speed is the difference between a meeting and a missed deal.")
add_table(doc,
    headers=["60-Day Success Benchmark", "Target", "Who Owns It"],
    rows=[
        ["Qualified leads per month",         "2–8",          "Marketing"],
        ["Lead response time",                "Under 24 hrs", "Sales"],
        ["Sales-accepted leads (SALs)",       "1+ per month", "Sales + Marketing"],
        ["Active late-stage opportunity",     "1 by Day 60",  "Sales"],
        ["ABM accounts touched per rep/week", "5+",           "Sales"],
    ],
    col_widths=[2.6, 1.4, 1.6]
)
add_body(doc,
    "If we hit these benchmarks, we expand the budget and launch LinkedIn outreach in Q3. "
    "If we miss, we adjust. Every decision is data-driven from here.", italic=True)
add_hr(doc)
add_body(doc,
    "April 6, 2026  ·  Chris, Head of Growth and Marketing  ·  SIGNiX",
    italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out = ("/Users/chris/Desktop/AI Summit 2026 Q2/"
       "proj-template-and-lease-SIGNiX-app/PROJECT-DOCS/"
       "SIGNiX_CRO_PaidMedia_Brief_4.6.26.docx")
doc.save(out)
print(f"Saved: {out}")
