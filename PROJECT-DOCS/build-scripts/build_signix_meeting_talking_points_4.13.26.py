#!/usr/bin/env python3
"""Build SIGNiX_MeetingTalkingPoints_4.13.26.docx — management meeting prep."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Brand colours ─────────────────────────────────────────────────────────────
GREEN = RGBColor(0x6d, 0xa3, 0x4a)
INK   = RGBColor(0x2e, 0x34, 0x40)
BODY  = RGBColor(0x54, 0x54, 0x54)
WHITE = RGBColor(0xff, 0xff, 0xff)
FONT  = "Calibri"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def style_run(run, color, bold=False, size=11, italic=False):
    run.font.name      = FONT
    run.font.color.rgb = color
    run.font.bold      = bold
    run.font.size      = Pt(size)
    run.font.italic    = italic

def _apply_inline(p, text, base_bold=False, base_italic=False, base_color=None):
    color = base_color or BODY
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            style_run(run, INK, bold=True, size=11)
        else:
            run = p.add_run(part)
            style_run(run, color, bold=base_bold, size=11, italic=base_italic)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=22)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=15)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "6da34a")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, GREEN, bold=True, size=12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)

def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    _apply_inline(p, text, base_italic=italic)
    p.paragraph_format.space_after = Pt(5)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after = Pt(3)
    _apply_inline(p, text)

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "2e3440")
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        style_run(run, WHITE, bold=True, size=10)
    for r_i, row in enumerate(rows):
        bg = "f8fafb" if r_i % 2 == 0 else "ffffff"
        for c_i, val in enumerate(row):
            c = tbl.rows[r_i + 1].cells[c_i]
            set_cell_bg(c, bg)
            c.paragraphs[0].clear()
            _apply_inline(c.paragraphs[0], val)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "6da34a")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    style_run(run, BODY, italic=True, size=11)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "d8dee9")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── BUILD ─────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(1.0)
sec.right_margin  = Inches(1.0)
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)

doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(11)

# ── Title ─────────────────────────────────────────────────────────────────────
add_h1(doc, "Management Meeting — Talking Points")

add_table(doc,
    headers=["Field", "Detail"],
    rows=[
        ("Presenter",  "Chris Teague, Head of Growth and Marketing"),
        ("Date",       "April 13, 2026"),
        ("Dashboards", "Marketing Dashboard · TruStage Portfolio Dashboard"),
        ("Order",      "Lead with Marketing. Close with TruStage."),
    ],
    col_widths=[1.4, 5.1]
)

add_blockquote(doc,
    "The marketing dashboard shows how we're bringing new leads in. "
    "The TruStage dashboard shows the revenue opportunity we already have."
)
add_hr(doc)

# ── Section 1: Marketing Dashboard ───────────────────────────────────────────
add_h2(doc, "Part 1 — Marketing Dashboard")
add_body(doc, "**Lead with:** \"We launched Google Ads on April 9th. Here's where we stand after the first month of data.\"")

add_h3(doc, "Point 1 — The top of the funnel is working")
add_bullet(doc, "$5,006 spent over 30 days (March 9 – April 7, 2026)")
add_bullet(doc, "1,224 clicks. 36 form submissions. Cost per lead: $139.")
add_bullet(doc, "Branded campaign outperforming RON: 35.6% CTR vs. 5.9% on RON")
add_bullet(doc, "Branded cost per lead: $48. RON cost per lead: $266.")

add_h3(doc, "Point 2 — We fixed three problems immediately")
add_bullet(doc, "**Off-hours spend fixed:** 31% of spend ($1,560) was running outside business hours. Now restricted to M–F, 9–5 EST.")
add_bullet(doc, "**Zero-conversion keywords paused:** 41 keywords burned $2,227 (44% of total spend) with zero leads. Paused.")
add_bullet(doc, "**Budget locked:** $100/day = $3,000/month. Was running higher. Now controlled.")
add_bullet(doc, "Net savings from these changes: ~$1,470/month redirected to performing keywords.")

add_h3(doc, "Point 3 — Lead follow-up is the open issue")
add_bullet(doc, "Average response time to a paid lead: **2.2 days**")
add_bullet(doc, "**0 out of 8 leads** contacted within 5 minutes")
add_bullet(doc, "1 out of 8 contacted within 10 minutes")
add_bullet(doc, "3 leads waited 24+ hours before any contact")
add_bullet(doc, "The ads are generating leads. The team needs to respond faster.")

add_h3(doc, "If asked about pipeline")
add_body(doc, "\"The pipeline section of the dashboard is pending. We need a HubSpot export by deal stage to populate it. That's the next data connection we're building.\"")

add_h3(doc, "Phase 2 — coming May 7")
add_bullet(doc, "Three new campaigns launch the week of May 7: Healthcare/Consent, Wealth Management, and Authentication/PKI")
add_bullet(doc, "Rep assignments confirmed: Aspen primary on all new campaigns")
add_hr(doc)

# ── Section 2: TruStage Dashboard ────────────────────────────────────────────
add_h2(doc, "Part 2 — TruStage Portfolio Dashboard")
add_body(doc, "**Lead with:** \"This is the depth-lever story. We already have these customers. The question is whether they're using authentication.\"")

add_h3(doc, "Point 1 — Transaction volume is growing")
add_bullet(doc, "14,087 total transactions tracked: January 2024 through March 2026")
add_bullet(doc, "2024: 4,698 transactions. 2025: 6,552 (+39% YoY). 2026 Q1: 2,837 (~945/month).")
add_bullet(doc, "Pace is accelerating. The base is healthy.")

add_h3(doc, "Point 2 — The authentication gap is real and quantifiable")
add_bullet(doc, "Overall auth rate across managed portfolio: **61.5%**")
add_bullet(doc, "~40% of transacting FIs use zero or near-zero authentication")
add_bullet(doc, "**Greenville Federal Bank:** 403 transactions. Zero authentications.")
add_bullet(doc, "**Commercial Bank of Texas:** 610 transactions. 0.7% auth rate.")
add_bullet(doc, "**First National Bank of Aspermont:** 289 transactions. 5.2% auth rate.")
add_bullet(doc, "No new customers needed. These are existing accounts leaving revenue on the table.")

add_h3(doc, "Point 3 — One account needs immediate attention")
add_bullet(doc, "**Old Monroe** was our flagship ID Verify account")
add_bullet(doc, "ID Verify usage has declined every month since December: 16 → 10 → 6 → 9 → 1")
add_bullet(doc, "This is a churn signal. It needs a CSM call this week.")

add_h3(doc, "If asked about channel mix")
add_body(doc, "\"API adoption is at 51% in Q1 2026. MyDoX is still at 49%. API integrations are where the deal thesis lives. We need that number to move.\"")

add_h3(doc, "If asked about revenue potential")
add_bullet(doc, "SMS auth: $0.25/event. ID Verify: $1.50/event.")
add_bullet(doc, "10 zero-auth accounts represent ~$9/month SMS potential if activated at baseline")
add_bullet(doc, "At scale, each activated account compounds — every transaction becomes a potential auth event")
add_hr(doc)

# ── Key Numbers Cheat Sheet ───────────────────────────────────────────────────
add_h2(doc, "Key Numbers — Quick Reference")
add_table(doc,
    headers=["Number", "What It Is", "Use It When..."],
    rows=[
        ("$139",       "Cost per lead (Google Ads)",             "Asked about paid media efficiency"),
        ("$48 vs $266","Branded CPL vs. RON CPL",                "Asked which campaign performs better"),
        ("36",         "Form submissions this period",           "Asked about lead volume"),
        ("2.2 days",   "Avg. lead response time",                "Discussing sales follow-up urgency"),
        ("0 / 8",      "Leads contacted within 5 minutes",       "Making the response time point"),
        ("$1,470/mo",  "Spend recaptured by pausing bad keywords","Showing proactive optimization"),
        ("14,087",     "Total TruStage transactions tracked",     "Setting the scale of the opportunity"),
        ("+39%",       "YoY transaction growth 2024 to 2025",    "Showing the base is healthy"),
        ("61.5%",      "Overall portfolio auth rate",            "Establishing auth baseline"),
        ("~40%",       "FIs at zero or near-zero auth",          "Making the depth-lever argument"),
        ("403 / 0",    "Greenville Federal: transactions / auths","Concrete zero-auth example"),
        ("610 / 0.7%", "Commercial Bank of TX: trans / auth rate","Second concrete zero-auth example"),
        ("51% API",    "API share of Q1 2026 transactions",      "If asked about channel mix"),
        ("May 7",      "Phase 2 Google Ads launch date",         "Showing what's coming next"),
    ],
    col_widths=[1.1, 2.6, 2.8]
)

add_hr(doc)
add_body(doc, "Prepared by Delford for Chris Teague · April 13, 2026 · SIGNiX Confidential", italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/chris/Desktop/AI Summit 2026 Q2/proj-template-and-lease-SIGNiX-app/PROJECT-DOCS/DELIVERABLES/SIGNiX_MeetingTalkingPoints_4.13.26.docx"
doc.save(out)
print(f"Saved: {out}")
