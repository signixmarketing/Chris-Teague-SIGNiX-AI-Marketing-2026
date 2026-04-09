#!/usr/bin/env python3
"""Build SIGNiX_GoogleAds_April2026.docx — full ad copy for five standalone campaigns."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Brand colours ─────────────────────────────────────────────────────────────
GREEN = RGBColor(0x6d, 0xa3, 0x4a)
INK   = RGBColor(0x2e, 0x34, 0x40)
BODY  = RGBColor(0x54, 0x54, 0x54)
MUTED = RGBColor(0x6b, 0x72, 0x80)
WHITE = RGBColor(0xff, 0xff, 0xff)
FONT  = "Calibri"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def style_run(run, color, bold=False, size=11, italic=False):
    run.font.name      = FONT
    run.font.color.rgb = color
    run.font.bold      = bold
    run.font.size      = Pt(size)
    run.font.italic    = italic

def _apply_inline(p, text, base_bold=False, base_italic=False, base_color=None):
    color = base_color or BODY
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            style_run(run, INK, bold=True, size=11)
        else:
            run = p.add_run(part)
            style_run(run, color, bold=base_bold, size=11, italic=base_italic)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=24)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=16)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "6da34a")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, GREEN, bold=True, size=12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    _apply_inline(p, text, base_italic=italic)
    p.paragraph_format.space_after = Pt(6)

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, MUTED, size=10, italic=True)
    p.paragraph_format.space_after = Pt(4)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    _apply_inline(p, text)

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "2e3440")
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        style_run(run, WHITE, bold=True, size=10)
    for r_i, row in enumerate(rows):
        bg_hex = "f8fafb" if r_i % 2 == 0 else "ffffff"
        for c_i, val in enumerate(row):
            c = tbl.rows[r_i + 1].cells[c_i]
            set_cell_bg(c, bg_hex)
            c.paragraphs[0].clear()
            _apply_inline(c.paragraphs[0], val)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "6da34a")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    style_run(run, BODY, italic=True, size=11)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "d8dee9")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── BUILD ─────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(1.0)
sec.right_margin  = Inches(1.0)
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)

doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(11)

# ── Title ─────────────────────────────────────────────────────────────────────
add_h1(doc, "SIGNiX Google Ads Copy — April 2026")

meta_rows = [
    ("Prepared by",   "Chris Teague, Head of Growth and Marketing"),
    ("Date",          "April 2026"),
    ("Status",        "Updated April 9, 2026 — 5 standalone campaigns, rep assignments confirmed by CRO Jesse"),
    ("Budget",        "~$3,000/month"),
    ("Landing page",  "signix.com/demo"),
    ("CTA",           "Request a Demo"),
    ("Ad format",     "Responsive Search Ads (RSA)"),
]
mt = doc.add_table(rows=len(meta_rows), cols=2)
mt.style = "Table Grid"
for i, (k, v) in enumerate(meta_rows):
    set_cell_bg(mt.rows[i].cells[0], "2e3440")
    mt.rows[i].cells[0].width = Inches(1.5)
    run_k = mt.rows[i].cells[0].paragraphs[0].add_run(k)
    style_run(run_k, WHITE, bold=True, size=10)
    set_cell_bg(mt.rows[i].cells[1], "f8fafb")
    run_v = mt.rows[i].cells[1].paragraphs[0].add_run(v)
    style_run(run_v, BODY, size=10)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_h3(doc, "Campaign structure")
add_body(doc, "Five standalone campaigns. Each campaign targets a distinct segment with its own keywords, fear-based headlines, and a rep reference in at least one description. All clicks go to **signix.com/demo**. Clean break from existing RON and Branded campaigns — baseline measurement starts at launch.")
add_bullet(doc, "**Phase 1 — Launch now:** Compliance / Legal (Aspen) + RON Institutional (Sam)")
add_bullet(doc, "**Phase 2 — Launch in 3–4 weeks:** Healthcare / Consent (Aspen + Chris), Wealth Management (Aspen), Authentication / PKI (Aspen)")
add_bullet(doc, "**Aspen:** primary rep on all new campaigns — Auth/PKI, Wealth Mgmt, Compliance/Legal, Healthcare/Consent")
add_bullet(doc, "**Sam:** RON Institutional + Debt/Legal (lighter load per CRO Jesse)")
add_bullet(doc, "**CSM team (2 people):** volume overflow backup for Aspen")
add_bullet(doc, "**Chris:** covers Healthcare/Consent leads when needed (medical devices / orthopedics background)")
add_note(doc, "Rep assignments confirmed by CRO Jesse, April 9, 2026. Budget path: reduce RON from $67 to $40/day, allocate $25–30/day to Phase 1 campaigns. Pending Chris confirmation.")
add_hr(doc)

# ── CAMPAIGN 1 ────────────────────────────────────────────────────────────────
add_h2(doc, "Campaign 1 — Authentication and Fraud Prevention")
add_body(doc, "**Sales rep:** Aspen Arias (owns Banks and Credit Unions per CRO Jesse) | **Phase:** 2 — launch in 3–4 weeks")
add_blockquote(doc, "Target: FI operations leaders and compliance officers at banks and credit unions who sign loan documents today but use zero authentication. Fear angle: a forged signature they can't disprove in court.")

add_h3(doc, "Keywords")
add_note(doc, "Use broad match modified or phrase match. Avoid exact match on head terms until data shows intent.")
add_table(doc,
    ["Keyword", "Match Type", "Intent"],
    [
        ("identity verification e-signature",       "Phrase", "High — active evaluation"),
        ("fraud prevention electronic signature",   "Phrase", "High — problem-aware"),
        ("authentication digital signature",        "Phrase", "High — solution-aware"),
        ("e-sign with identity verification",       "Phrase", "High — active evaluation"),
        ("secure document signing financial",       "Broad",  "Medium — research phase"),
        ("multi-factor authentication document",    "Broad",  "Medium — IT buyer"),
        ("loan document fraud prevention",          "Phrase", "High — pain-aware"),
        ("who signed the document proof",           "Broad",  "Medium — post-incident"),
        ("biometric e-signature",                   "Phrase", "Medium — feature search"),
        ("KBA knowledge based authentication",      "Phrase", "High — compliance-aware"),
        ("digital signatures pki",                  "Phrase", "Medium — technical/developer buyer"),
    ],
    col_widths=[3.0, 1.2, 2.3]
)

add_h3(doc, "Headlines")
add_note(doc, "Google rotates up to 15 headlines. Max 30 characters each. Pinned positions noted where intent is strong.")
add_table(doc,
    ["#", "Headline", "Chars", "Pin?"],
    [
        ("1",  "Prove Who Signed It",              "19", "Pin 1 — lead with the fear hook"),
        ("2",  "Stop Loan Document Fraud",          "24", ""),
        ("3",  "Sign With Identity Proof",          "24", ""),
        ("4",  "SIGNiX: ID at Every Signing",       "27", ""),
        ("5",  "Authentication Built In",           "22", ""),
        ("6",  "Digital Signatures + ID Check",     "29", ""),
        ("7",  "Fraud-Proof Your Documents",        "26", ""),
        ("8",  "Who Really Signed That Doc?",       "27", ""),
        ("9",  "Sign. Verify. Defend.",              "21", ""),
        ("10", "ID Verification at Signing",        "26", ""),
        ("11", "Request a Demo Today",              "20", "Pin 3 — close with CTA"),
        ("12", "Talk to Aspen at SIGNiX",           "23", ""),
        ("13", "Protect Every Transaction",         "25", ""),
        ("14", "No More Disputed Signatures",       "27", ""),
        ("15", "Flex API + Authentication",         "25", ""),
    ],
    col_widths=[0.3, 2.7, 0.6, 2.9]
)

add_h3(doc, "Descriptions")
add_note(doc, "Google shows up to 2 descriptions per ad. Max 90 characters each.")
add_table(doc,
    ["#", "Description", "Chars"],
    [
        ("1", "SIGNiX verifies who signs every document. Stop disputes before they start. Talk to Aspen.",        "90"),
        ("2", "When a member disputes a signature, you need proof. SIGNiX attaches ID to every sign.",           "86"),
        ("3", "Fraud starts where identity verification ends. SIGNiX closes that gap. Request a demo.",          "87"),
        ("4", "Your e-sign has a signature. Does it have ID proof? SIGNiX does. Request a demo.",                "81"),
    ],
    col_widths=[0.3, 5.7, 0.5]
)
add_hr(doc)

# ── CAMPAIGN 2 ────────────────────────────────────────────────────────────────
add_h2(doc, "Campaign 2 — Wealth Management")
add_body(doc, "**Sales rep:** Aspen Arias (confirmed by CRO Jesse, April 9, 2026) | **Phase:** 2 — launch in 3–4 weeks")
add_blockquote(doc, "Target: RIAs, broker-dealers, and financial advisors who use standard e-sign for client agreements. Fear angle: a client dispute or FINRA exam that exposes a signature they cannot legally prove.")

add_h3(doc, "Keywords")
add_note(doc, "Wealth management searches tend to be compliance-driven. Bid on regulatory terms and role-specific phrases.")
add_table(doc,
    ["Keyword", "Match Type", "Intent"],
    [
        ("FINRA compliant e-signature",                 "Phrase", "High — compliance-driven"),
        ("wealth management digital signature",         "Phrase", "High — role + solution"),
        ("financial advisor document signing",          "Phrase", "High — workflow match"),
        ("RIA electronic signature solution",           "Phrase", "High — active evaluation"),
        ("broker dealer e-signature",                   "Phrase", "High — role-specific"),
        ("SEC compliant document signing",              "Phrase", "Medium — regulatory"),
        ("registered investment advisor e-sign",        "Broad",  "Medium — research phase"),
        ("client agreement digital signature",          "Phrase", "High — feature match"),
        ("investment advisory document signing",        "Broad",  "Medium — workflow pain"),
        ("e-signature financial services compliance",   "Broad",  "Medium — IT buyer"),
        ("power of attorney e-signature",               "Phrase", "High — high-stakes authority doc"),
    ],
    col_widths=[3.0, 1.2, 2.3]
)

add_h3(doc, "Headlines")
add_table(doc,
    ["#", "Headline", "Chars", "Pin?"],
    [
        ("1",  "FINRA-Ready Digital Signatures",  "30", "Pin 1 — lead with compliance"),
        ("2",  "Defend Every Client Document",    "28", ""),
        ("3",  "E-Sign Built for RIAs",           "21", ""),
        ("4",  "Audit-Ready Client Signatures",   "29", ""),
        ("5",  "Prove Identity at Each Signing",  "30", ""),
        ("6",  "Stop Signature Disputes Now",     "27", ""),
        ("7",  "Talk to a SIGNiX Specialist",      "27", ""),
        ("8",  "Request a Demo Today",            "20", "Pin 3 — close with CTA"),
        ("9",  "PKI Signatures for Advisors",     "27", ""),
        ("10", "SIGNiX for Broker-Dealers",       "25", ""),
        ("11", "Beyond Standard E-Signature",     "27", ""),
        ("12", "Identity Proof on Every Doc",     "27", ""),
        ("13", "Digital Signatures With ID",      "25", ""),
        ("14", "Client Docs That Hold Up",        "23", ""),
        ("15", "SIGNiX Flex API",                 "15", ""),
    ],
    col_widths=[0.3, 2.7, 0.6, 2.9]
)

add_h3(doc, "Descriptions")
add_table(doc,
    ["#", "Description", "Chars"],
    [
        ("1", "A disputed client document can cost more than the transaction. SIGNiX proves who signed.",          "88"),
        ("2", "Standard e-sign proves a click happened. SIGNiX proves who the person was. Request a demo.",       "88"),
        ("3", "PKI-based digital signatures attach verified identity to every client document you send.",          "88"),
        ("4", "Your client deserves a signature that holds up in court. Request a demo at signix.com/demo.",       "89"),
    ],
    col_widths=[0.3, 5.7, 0.5]
)
add_hr(doc)

# ── CAMPAIGN 3 ────────────────────────────────────────────────────────────────
add_h2(doc, "Campaign 3 — RON Institutional (Phase 1)")
add_body(doc, "**Sales rep:** Sam West (owns Debt Collection, Law Firms, and RON per CRO Jesse) | **Phase:** 1 — launch now")
add_blockquote(doc, "Target: debt buyers, law firms, legal publishers, and process servers that need Remote Online Notary for documents that must hold up in court. Fear angle: a legal notice that gets thrown out because it was never properly notarized.")

add_h3(doc, "Keywords")
add_note(doc, "RON searches are specific. Bid aggressively on exact terms — competition from DocuSign and Notarize.com is real but not dominant in niche legal/debt verticals.")
add_table(doc,
    ["Keyword", "Match Type", "Intent"],
    [
        ("remote online notary service",            "Phrase", "High — active evaluation"),
        ("RON notarization platform",               "Phrase", "High — solution-aware"),
        ("electronic notarization service",         "Phrase", "High — solution-aware"),
        ("notarize documents online",               "Phrase", "High — ready to buy"),
        ("debt collection e-signature",             "Phrase", "High — role-specific"),
        ("legal notice notarization",               "Phrase", "High — workflow pain"),
        ("process server electronic signature",     "Phrase", "High — role-specific"),
        ("legal publisher document signing",        "Broad",  "Medium — niche match"),
        ("RON API integration",                     "Phrase", "High — developer/ISV"),
        ("digital notary online compliance",        "Broad",  "Medium — compliance"),
        ("e notary platforms",                      "Phrase", "Low comp, growing — institutional buyer language"),
        ("electronic notarization platform",        "Phrase", "Emerging term, near-zero competition"),
    ],
    col_widths=[3.0, 1.2, 2.3]
)

add_h3(doc, "Headlines")
add_table(doc,
    ["#", "Headline", "Chars", "Pin?"],
    [
        ("1",  "Remote Online Notary Ready",      "26", "Pin 1 — match search intent"),
        ("2",  "RON for Legal Publishers",        "24", ""),
        ("3",  "Notarize Documents Online",       "25", ""),
        ("4",  "SIGNiX RON: State Compliant",     "27", ""),
        ("5",  "Legal Notices Need RON",          "22", ""),
        ("6",  "Debt Collection E-Sign",          "22", ""),
        ("7",  "DocuSign Can't Do RON",           "21", ""),
        ("8",  "Sign and Notarize in One Step",   "29", ""),
        ("9",  "Talk to Sam at SIGNiX",            "21", ""),
        ("10", "Request a Demo Today",            "20", "Pin 3 — close with CTA"),
        ("11", "RON Built for Compliance",        "24", ""),
        ("12", "Court-Admissible Signatures",     "27", ""),
        ("13", "SIGNiX Flex API",                 "15", ""),
        ("14", "E-Sign + RON in One Platform",    "28", ""),
        ("15", "State-Approved Notarization",     "27", ""),
    ],
    col_widths=[0.3, 2.7, 0.6, 2.9]
)

add_h3(doc, "Descriptions")
add_table(doc,
    ["#", "Description", "Chars"],
    [
        ("1", "State law requires notarization on many legal notices. DocuSign can't do RON. SIGNiX can.",        "89"),
        ("2", "SIGNiX handles Remote Online Notary and e-sign in one platform. Talk to Sam.",                    "77"),
        ("3", "Debt buyers and legal publishers trust SIGNiX for documents that hold up in court.",              "82"),
        ("4", "One platform for e-sign, ID verification, and Remote Online Notary. Request a demo.",             "83"),
    ],
    col_widths=[0.3, 5.7, 0.5]
)
add_hr(doc)

# ── CAMPAIGN 4 ────────────────────────────────────────────────────────────────
add_h2(doc, "Campaign 4 — Compliance / Legal (Phase 1)")
add_body(doc, "**Sales rep:** Aspen Arias | **Phase:** 1 — launch now")
add_blockquote(doc, "Target: legal ops managers, contract teams, and compliance officers researching whether electronic or digital signatures will hold up in court. Fear angle: signing a contract with a method that gets thrown out in a dispute.")

add_h3(doc, "Keywords")
add_note(doc, "Start all as phrase match. 'electronic signature verification statement example' is a sleeper — +1,500% YoY, near-zero competition. Consider broad match to capture the surge.")
add_table(doc,
    ["Keyword", "Match Type", "Intent"],
    [
        ("are electronic signatures legally binding",          "Phrase", "High — active research, ready to evaluate"),
        ("are digital signatures legally binding",             "Phrase", "High — fastest 3-mo growth in bucket, +57%"),
        ("signatures on contracts",                            "Phrase", "Medium — broad but low comp, top-of-funnel"),
        ("legal document signing",                             "Phrase", "Medium — workflow-specific, legal/debt sector"),
        ("electronic signature verification statement example","Broad",  "High — surging +1,500% YoY, near-zero comp"),
    ],
    col_widths=[3.0, 1.2, 2.3]
)

add_h3(doc, "Headlines")
add_note(doc, "Google rotates up to 15 headlines. Max 30 characters each.")
add_table(doc,
    ["#", "Headline", "Chars", "Pin?"],
    [
        ("1",  "Court-Ready Digital Signatures",  "30", "Pin 1 — lead with authority"),
        ("2",  "Prove Your Signature Is Legal",   "29", ""),
        ("3",  "Legal Digital Signatures",        "24", ""),
        ("4",  "PKI Beats Standard E-Sign",       "25", ""),
        ("5",  "Sign. Prove. Defend.",             "20", ""),
        ("6",  "More Than a Click. Proven ID.",   "29", ""),
        ("7",  "Stop Disputed Signatures",        "24", ""),
        ("8",  "Signatures With Identity Proof",  "30", ""),
        ("9",  "Is Your E-Sign Legally Sound?",   "29", ""),
        ("10", "SIGNiX: Legal Signature Proof",   "29", ""),
        ("11", "Request a Demo Today",            "20", "Pin 3 — CTA"),
        ("12", "Talk to Aspen at SIGNiX",         "23", ""),
        ("13", "Legal Docs That Stand Up",        "24", ""),
        ("14", "Authentication + Digital Sign",   "29", ""),
        ("15", "SIGNiX Flex API",                 "15", ""),
    ],
    col_widths=[0.3, 2.7, 0.6, 2.9]
)

add_h3(doc, "Descriptions")
add_note(doc, "Google shows up to 2 descriptions per ad. Max 90 characters each.")
add_table(doc,
    ["#", "Description", "Chars"],
    [
        ("1", "Standard e-sign proves a click happened. SIGNiX proves who signed. Ask Aspen how.",         "82"),
        ("2", "A signature that can't be proven in court isn't worth much. SIGNiX changes that.",          "81"),
        ("3", "PKI-based digital signatures create legal proof at every step. Request a demo.",             "78"),
        ("4", "Your contracts need more than an image of a name. SIGNiX adds verified identity.",          "80"),
    ],
    col_widths=[0.3, 5.7, 0.5]
)
add_hr(doc)

# ── CAMPAIGN 5 ────────────────────────────────────────────────────────────────
add_h2(doc, "Campaign 5 — Healthcare / Consent (Phase 2)")
add_body(doc, "**Sales rep:** Aspen Arias (primary) + Chris Teague when needed (orthopedics / medical device background) | **Phase:** 2 — launch in 3–4 weeks")
add_blockquote(doc, "Target: healthcare organizations, patient intake teams, and HIPAA compliance officers handling consent forms. Fear angle: a consent dispute that exposes a form you can't verify was signed by the right person.")

add_h3(doc, "Keywords")
add_note(doc, "HIPAA buyers are cautious and compliance-driven. Pair regulatory language with workflow-specific terms.")
add_table(doc,
    ["Keyword", "Match Type", "Intent"],
    [
        ("hipaa electronic signature",   "Phrase", "High — compliance-driven, medium competition"),
        ("e signature consent form",     "Phrase", "High — growing +67% 3-mo, consent = authority doc"),
        ("e signature consent",          "Phrase", "Medium — stable, pairs well with consent form term"),
    ],
    col_widths=[3.0, 1.2, 2.3]
)

add_h3(doc, "Headlines")
add_table(doc,
    ["#", "Headline", "Chars", "Pin?"],
    [
        ("1",  "HIPAA-Ready E-Signatures",         "24", "Pin 1 — match compliance intent"),
        ("2",  "Patient Consent With ID Proof",    "29", ""),
        ("3",  "Verified Consent Forms",           "22", ""),
        ("4",  "HIPAA Compliant Digital Sign",     "27", ""),
        ("5",  "Protect Patient Consent Docs",     "28", ""),
        ("6",  "E-Sign Built for Healthcare",      "27", ""),
        ("7",  "Consent Forms That Hold Up",       "25", ""),
        ("8",  "Stop Consent Disputes Now",        "25", ""),
        ("9",  "Signed, Verified, Defensible",     "28", ""),
        ("10", "Request a Demo Today",             "20", "Pin 3 — CTA"),
        ("11", "Talk to Aspen at SIGNiX",          "23", ""),
        ("12", "Healthcare E-Sign + ID Check",     "27", ""),
        ("13", "Prove Who Signed the Consent",     "28", ""),
        ("14", "Beyond HIPAA. Proven Identity.",   "30", ""),
        ("15", "SIGNiX Flex API",                  "15", ""),
    ],
    col_widths=[0.3, 2.7, 0.6, 2.9]
)

add_h3(doc, "Descriptions")
add_table(doc,
    ["#", "Description", "Chars"],
    [
        ("1", "A patient dispute starts with a signature you can't verify. SIGNiX closes that gap.",       "83"),
        ("2", "HIPAA compliance needs more than a checkbox. SIGNiX adds verified ID to consent forms.",    "85"),
        ("3", "SIGNiX attaches verified identity to every consent form. Request a demo today.",            "78"),
        ("4", "Consent disputes cost more than the fix. SIGNiX proves who signed every document.",        "81"),
    ],
    col_widths=[0.3, 5.7, 0.5]
)
add_hr(doc)

# ── Notes ─────────────────────────────────────────────────────────────────────
add_h2(doc, "Launch Notes")

add_h3(doc, "How Google uses RSA copy")
add_body(doc, "Google assembles each ad from your headlines and descriptions automatically. It tests combinations and favors the highest-performing ones. Providing 15 headlines and 4 descriptions gives Google maximum flexibility.")
add_body(doc, "**Pinning** a headline locks it to a specific position (1, 2, or 3) on every impression. Use pins sparingly. Pinning too many headlines reduces Google's ability to optimize.")

add_h3(doc, "Conversion tracking")
add_body(doc, "Every click goes to **signix.com/demo**. Make sure the demo form submission fires a Google Ads conversion event. Without conversion tracking, Smart Bidding has no signal to optimize against.")
add_bullet(doc, "Tag the thank-you page or the form submission event in Google Tag Manager")
add_bullet(doc, "Set conversion goal: Demo Request (lead)")
add_bullet(doc, "Import conversion data into HubSpot for closed-loop reporting")

add_h3(doc, "Negative keywords (recommended)")
add_table(doc,
    ["Negative keyword", "Reason"],
    [
        ("free",                    "Filters out consumers looking for no-cost tools"),
        ("individual notary",       "We are not targeting solo notary professionals"),
        ("notary public near me",   "Local intent — wrong audience"),
        ("HelloSign",               "Competitor brand — separate campaign needed if desired"),
        ("DocuSign free trial",     "Competitor brand + free intent"),
        ("jobs",                    "Filters job seekers searching e-sign roles"),
        ("template",                "Filters DIY document template searches"),
        ("how to notarize",         "Informational intent — not ready to buy"),
    ],
    col_widths=[2.5, 4.0]
)

add_h3(doc, "Rep name usage")
add_body(doc, "Sam and Aspen are referenced by first name in descriptions. This is intentional. It makes the ad feel less like a brand campaign and more like a direct conversation. The name signals there is a real person ready to respond.")
add_body(doc, "If sales leadership prefers not to use rep names in ads, the fallback is: **Talk to a SIGNiX specialist today.**")

add_hr(doc)
add_note(doc, "Source of truth: PROJECT-DOCS/build-scripts/build_signix_google_ads_docx.py | Updated April 9, 2026 — 5 campaigns, 12 validated keywords, rep assignments confirmed by CRO Jesse")

# ── Save ──────────────────────────────────────────────────────────────────────
import shutil, os

OUT = "PROJECT-DOCS/SIGNiX_GoogleAds_4.9.26.docx"
doc.save(OUT)
print(f"Saved: {OUT}")

DESKTOP = os.path.expanduser("~/Desktop/SIGNiX_GoogleAds_4.9.26.docx")
shutil.copy2(OUT, DESKTOP)
print(f"Copied to desktop: {DESKTOP}")
