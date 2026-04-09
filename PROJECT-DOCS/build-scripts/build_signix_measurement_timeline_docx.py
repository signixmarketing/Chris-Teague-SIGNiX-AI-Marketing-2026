#!/usr/bin/env python3
"""Build SIGNiX_CampaignMeasurement_April2026.docx — 90-day measurement timeline."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, shutil, os

# ── Brand colours ─────────────────────────────────────────────────────────────
GREEN  = RGBColor(0x6d, 0xa3, 0x4a)
INK    = RGBColor(0x2e, 0x34, 0x40)
BODY   = RGBColor(0x54, 0x54, 0x54)
MUTED  = RGBColor(0x6b, 0x72, 0x80)
WHITE  = RGBColor(0xff, 0xff, 0xff)
FONT   = "Calibri"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def style_run(run, color, bold=False, size=11, italic=False):
    run.font.name      = FONT
    run.font.color.rgb = color
    run.font.bold      = bold
    run.font.size      = Pt(size)
    run.font.italic    = italic

def _apply_inline(p, text, base_bold=False, base_italic=False, base_color=None):
    color = base_color or BODY
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            style_run(run, INK, bold=True, size=11)
        else:
            run = p.add_run(part)
            style_run(run, color, bold=base_bold, size=11, italic=base_italic)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=22)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "6da34a")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_h3(doc, text, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, color or GREEN, bold=True, size=12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)

def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    _apply_inline(p, text, base_italic=italic)
    p.paragraph_format.space_after = Pt(5)

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, MUTED, size=10, italic=True)
    p.paragraph_format.space_after = Pt(4)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    _apply_inline(p, text)

def add_table(doc, headers, rows, col_widths=None, header_bg="2e3440"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, header_bg)
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        style_run(run, WHITE, bold=True, size=10)
    for r_i, row in enumerate(rows):
        bg_hex = "f8fafb" if r_i % 2 == 0 else "ffffff"
        for c_i, val in enumerate(row):
            c = tbl.rows[r_i + 1].cells[c_i]
            set_cell_bg(c, bg_hex)
            c.paragraphs[0].clear()
            _apply_inline(c.paragraphs[0], val)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "6da34a")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    style_run(run, BODY, italic=True, size=11)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "d8dee9")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── BUILD ─────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(1.0)
sec.right_margin  = Inches(1.0)
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)

doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(11)

# ── Title ─────────────────────────────────────────────────────────────────────
add_h1(doc, "SIGNiX Google Ads — 90-Day Measurement Timeline")

meta_rows = [
    ("Prepared by",    "Chris Teague, Head of Growth and Marketing"),
    ("Date",           "April 2026"),
    ("Campaign",       "Three audience ad groups — Authentication/Fraud, Wealth Management, Debt/Legal/RON"),
    ("Budget",         "~$3,000/month"),
    ("Sales cycle",    "60 to 90 days average close time"),
    ("Audience",       "CRO, CEO"),
]
mt = doc.add_table(rows=len(meta_rows), cols=2)
mt.style = "Table Grid"
for i, (k, v) in enumerate(meta_rows):
    set_cell_bg(mt.rows[i].cells[0], "2e3440")
    mt.rows[i].cells[0].width = Inches(1.5)
    run_k = mt.rows[i].cells[0].paragraphs[0].add_run(k)
    style_run(run_k, WHITE, bold=True, size=10)
    set_cell_bg(mt.rows[i].cells[1], "f8fafb")
    run_v = mt.rows[i].cells[1].paragraphs[0].add_run(v)
    style_run(run_v, BODY, size=10)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_blockquote(doc, "Our average sales cycle is 60 to 90 days. Judging this campaign on closed revenue at 30 days is the wrong measure. This document sets the right benchmarks at each stage so we protect the campaign and make informed decisions.")

add_hr(doc)

# ── Why the timeline matters ───────────────────────────────────────────────────
add_h2(doc, "Why the Sales Cycle Changes How We Measure")
add_body(doc, "A prospect who fills out our demo form today will not close until June or July at the earliest. That is not a campaign failure. That is our sales cycle working normally.")
add_body(doc, "If we measure only closed deals at 30 days, we will always see zero. That creates pressure to pull budget from a campaign that may already be filling the pipeline. This framework prevents that.")
add_body(doc, "**Each phase has its own success signal.** Meet the phase benchmark and the campaign is on track, even before a deal closes.")

add_hr(doc)

# ── Phase 1 ────────────────────────────────────────────────────────────────────
add_h2(doc, "Phase 1 — Days 1 to 30: Is the Targeting Working?")
add_h3(doc, "What we are measuring")
add_body(doc, "This phase answers one question: are the right people clicking? We are not measuring revenue. We are measuring signal quality.")

add_table(doc,
    ["Metric", "Target", "What It Tells Us"],
    [
        ("Impressions",              "5,000+ per month",    "Ads are showing to the right search queries"),
        ("Click-through rate (CTR)", "2% or higher",        "Headlines are resonating with the audience"),
        ("Cost per click (CPC)",     "Under $8",            "Budget is being used efficiently"),
        ("Demo form submissions",    "10+ per month",       "Landing page and CTA are converting"),
        ("Cost per lead (CPL)",      "Under $300",          "Acquisition cost is sustainable vs. deal value"),
    ],
    col_widths=[2.2, 1.6, 2.7]
)

add_h3(doc, "What to report to leadership at Day 30")
add_body(doc, "Share CTR, total form fills, and cost per lead. Frame it this way: **'We have X qualified leads in the pipeline at $Y per lead. First deals are expected to close in June based on our 60 to 90 day cycle.'**")
add_note(doc, "Do not report on closed revenue at Day 30. It is not yet possible and sets a false expectation.")

add_hr(doc)

# ── Phase 2 ────────────────────────────────────────────────────────────────────
add_h2(doc, "Phase 2 — Days 30 to 60: Is the Pipeline Building?")
add_h3(doc, "What we are measuring")
add_body(doc, "This phase answers: are leads turning into real conversations? Sales activity in HubSpot is the primary data source here.")

add_table(doc,
    ["Metric", "Target", "What It Tells Us"],
    [
        ("Leads contacted by sales",        "80% of form fills within 48 hours", "Sales is following up on paid leads"),
        ("Qualified conversations",         "5+ active evaluations",              "Leads have real intent, not just curiosity"),
        ("Pipeline value influenced",       "$50,000+",                           "Campaign is building measurable revenue potential"),
        ("Demo-to-conversation rate",       "50% or higher",                      "Demo form is attracting the right buyers"),
        ("Accounts in active evaluation",   "3 or more",                          "Sales cycle is progressing normally"),
    ],
    col_widths=[2.5, 2.0, 2.0]
)

add_h3(doc, "What to report to leadership at Day 60")
add_body(doc, "Share pipeline value and active account count. Frame it: **'We have $X in influenced pipeline across Y accounts. Based on our close rate, we expect Z deals to close by [date].'**")
add_note(doc, "If HubSpot is not logging lead source and sales activity consistently, pipeline reporting will be incomplete. Flag this to the CRO now, not at Day 60.")

add_hr(doc)

# ── Phase 3 ────────────────────────────────────────────────────────────────────
add_h2(doc, "Phase 3 — Days 60 to 90: Is the Campaign Generating Revenue?")
add_h3(doc, "What we are measuring")
add_body(doc, "This is the first phase where closed revenue is a fair measure. Deals from Day 1 form fills are now old enough to have closed or advanced to final stage.")

add_table(doc,
    ["Metric", "Target", "What It Tells Us"],
    [
        ("Closed deals from campaign leads",      "2 or more",           "Campaign is generating revenue"),
        ("Revenue closed or contracted",          "$40,000+",            "Campaign ROI is positive vs. $9K spend"),
        ("Deals in final stage (not yet closed)", "3 or more",           "Pipeline health for next 30 days"),
        ("Cost per closed deal",                  "Under $4,500",        "Acquisition cost vs. deal value is sound"),
        ("Authentication activations",            "1 FI pilot active",   "Depth strategy is moving"),
    ],
    col_widths=[2.8, 1.5, 2.2]
)

add_h3(doc, "What to report to leadership at Day 90")
add_body(doc, "This is the full ROI report. Share closed revenue, pipeline value, cost per deal, and a recommendation: continue, adjust, or expand. **At $3,000/month, two closed deals in 90 days with any meaningful transaction volume pays for the campaign many times over.**")

add_hr(doc)

# ── ROI frame ─────────────────────────────────────────────────────────────────
add_h2(doc, "ROI Reference Frame")
add_body(doc, "Use this to anchor leadership conversations at any phase.")

add_table(doc,
    ["Scenario", "Deals Closed", "Est. Annual Transaction Value", "Campaign Cost (90 days)", "ROI"],
    [
        ("Conservative", "1",  "$20,000",  "$9,000", "2.2x"),
        ("Base case",    "2",  "$50,000",  "$9,000", "5.6x"),
        ("Strong",       "3+", "$100,000+","$9,000", "11x+"),
    ],
    col_widths=[1.5, 1.1, 2.2, 1.8, 0.9]
)
add_note(doc, "Transaction value estimates based on average FI deal size. Adjust with CRO input once first deals close.")

add_hr(doc)

# ── Decision rules ─────────────────────────────────────────────────────────────
add_h2(doc, "Decision Rules — When to Adjust vs. Stay the Course")
add_table(doc,
    ["Signal", "Action"],
    [
        ("CTR below 1% at Day 30",                         "Test new headlines. Do not cut budget yet."),
        ("CPL above $500 at Day 30",                       "Tighten keyword match types. Review negative keyword list."),
        ("Fewer than 5 form fills at Day 30",              "Check landing page — may be the conversion bottleneck, not the ads."),
        ("Sales not following up within 48 hours",         "Escalate to CRO. A lead problem is not an ad problem."),
        ("Zero qualified conversations at Day 60",         "Review lead quality. Consider tightening audience targeting."),
        ("Pipeline value under $20K at Day 60",            "Assess whether ICP is right. May need segment adjustment."),
        ("2 or more deals closed or in final stage at 90", "Expand budget. The model is working."),
    ],
    col_widths=[3.2, 3.3]
)

add_hr(doc)
add_note(doc, "Source of truth: PROJECT-DOCS/build-scripts/build_signix_measurement_timeline_docx.py | Generated April 2026")

# ── Save ──────────────────────────────────────────────────────────────────────
OUT     = "PROJECT-DOCS/SIGNiX_CampaignMeasurement_April2026.docx"
DESKTOP = os.path.expanduser("~/Desktop/SIGNiX_CampaignMeasurement_April2026.docx")

doc.save(OUT)
print(f"Saved: {OUT}")
shutil.copy2(OUT, DESKTOP)
print(f"Copied to desktop: {DESKTOP}")
