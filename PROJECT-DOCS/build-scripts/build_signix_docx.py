#!/usr/bin/env python3
"""Build SIGNIX-PAID-MEDIA-PLAN-2026.docx with SIGNiX branding."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Brand colours ─────────────────────────────────────────────────────────────
GREEN = RGBColor(0x6d, 0xa3, 0x4a)
INK   = RGBColor(0x2e, 0x34, 0x40)
BODY  = RGBColor(0x54, 0x54, 0x54)
WHITE = RGBColor(0xff, 0xff, 0xff)
FONT  = "Calibri"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def style_run(run, color, bold=False, size=11, italic=False):
    run.font.name      = FONT
    run.font.color.rgb = color
    run.font.bold      = bold
    run.font.size      = Pt(size)
    run.font.italic    = italic

def _apply_inline(p, text, base_bold=False, base_italic=False, base_color=None):
    """Render **bold** inline markdown within a paragraph."""
    color = base_color or BODY
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            style_run(run, INK, bold=True, size=11)
        else:
            run = p.add_run(part)
            style_run(run, color, bold=base_bold, size=11, italic=base_italic)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=24)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=16)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # green bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "6da34a")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, GREEN, bold=True, size=12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    _apply_inline(p, text, base_italic=italic)
    p.paragraph_format.space_after = Pt(6)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    _apply_inline(p, text)

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # header row
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "2e3440")
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        style_run(run, WHITE, bold=True, size=10)
    # data rows
    for r_i, row in enumerate(rows):
        bg = "f8fafb" if r_i % 2 == 0 else "ffffff"
        for c_i, val in enumerate(row):
            c = tbl.rows[r_i + 1].cells[c_i]
            set_cell_bg(c, bg)
            c.paragraphs[0].clear()
            _apply_inline(c.paragraphs[0], val)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
    # column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "6da34a")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    style_run(run, BODY, italic=True, size=11)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "d8dee9")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── BUILD ─────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(1.0)
sec.right_margin  = Inches(1.0)
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)

# Default paragraph font
doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(11)

# ── Title ─────────────────────────────────────────────────────────────────────
add_h1(doc, "SIGNiX Paid Media Plan — 2026")

# Meta info table
meta_rows = [
    ("Prepared by", "Chris, Head of Growth and Marketing"),
    ("Date",        "April 2026"),
    ("Status",      "Proposed — pending CEO approval"),
    ("Budget",      "$5,000/month (reallocation of existing spend)"),
    ("Objective",   "Shift paid media investment from individual notary leads to high-value B2B pipeline — FI operations and IT buyers evaluating e-sign and authentication, and ISV developers evaluating the Flex API."),
]
mt = doc.add_table(rows=len(meta_rows), cols=2)
mt.style = "Table Grid"
for i, (k, v) in enumerate(meta_rows):
    set_cell_bg(mt.rows[i].cells[0], "2e3440")
    mt.rows[i].cells[0].width = Inches(1.4)
    run_k = mt.rows[i].cells[0].paragraphs[0].add_run(k)
    style_run(run_k, WHITE, bold=True, size=10)
    set_cell_bg(mt.rows[i].cells[1], "f8fafb")
    run_v = mt.rows[i].cells[1].paragraphs[0].add_run(v)
    style_run(run_v, BODY, size=10)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_h3(doc, "Related Documents")
add_bullet(doc, "Full growth strategy: SIGNIX-GROWTH-STRATEGY.md")
add_bullet(doc, "Three-track marketing plan: SIGNIX-MARKETING-PLAN-2026.md")
add_bullet(doc, "Executive summary (one-page): SIGNIX-PAID-MEDIA-EXEC-SUMMARY.md")
add_bullet(doc, "Slide deck: SIGNIX-PAID-MEDIA-SLIDES.html")
add_hr(doc)

# ── Section 1 ─────────────────────────────────────────────────────────────────
add_h2(doc, "1. The Case for Reallocation")
add_h3(doc, "Current state")
add_body(doc, "SIGNiX's paid media spend (~$5,000/month) is concentrated on Google Ads targeting individual notary and RON queries. These campaigns attract individual notary professionals — low ACV, high-touch customers who require a full sales close plus dedicated CSM onboarding.")
add_body(doc, "**The economics are broken at the unit level:**")
add_body(doc, "When the salary cost of a sales representative and Client Success Manager (Gina) is divided by the percentage of time spent per notary lead — from first touch through closed deal and full onboarding — the contribution margin per notary customer is likely near zero or negative.")
add_body(doc, "**Key signals:**")
add_bullet(doc, "Individual notary deals: low ACV, one-time onboarding cost, low recurring transaction volume")
add_bullet(doc, "Staff time per deal: too high relative to deal value")
add_bullet(doc, "Scalability: none — each deal requires the same manual motion regardless of volume")

add_h3(doc, "Proposed state")
add_body(doc, "Redirect the same $5,000/month toward audiences that generate high-value, recurring-revenue customers:")
add_bullet(doc, "**Financial institutions (FIs):** Banks, credit unions, and lenders evaluating e-sign and authentication for lending, compliance, and operational workflows. Higher ACV, recurring transaction revenue per account.")
add_bullet(doc, "**ISV developers and platform partners:** Product managers and engineers at LOS, fintech, and banking software companies evaluating SIGNiX's Flex API for embedded signing. Single integration multiplies the transaction base.")
add_body(doc, "**The shift does not require a budget increase.** It requires redirecting existing spend toward higher-leverage audiences.")
add_hr(doc)

# ── Section 2 ─────────────────────────────────────────────────────────────────
add_h2(doc, "2. Budget Allocation")
add_table(doc,
    headers=["Channel", "Monthly Allocation", "Rationale"],
    rows=[
        ["Google Search Ads", "$3,000", "Capture active evaluators searching for e-sign and authentication solutions; competitor displacement"],
        ["LinkedIn Sponsored Content", "$1,500", "Reach cold but qualified FI and ISV audiences not actively searching; builds brand with the right personas"],
        ["Test / reserve budget", "$500", "A/B test landing pages; pilot new keyword experiments; buffer for CPC volatility"],
        ["**Total**", "**$5,000**", "No increase from current spend"],
    ],
    col_widths=[2.0, 1.5, 3.0]
)
add_hr(doc)

# ── Section 3 ─────────────────────────────────────────────────────────────────
add_h2(doc, "3. Google Search Ads Strategy ($3,000/month)")
add_h3(doc, "Strategic principle")
add_body(doc, "Avoid head-to-head competition on generic e-sign terms where DocuSign and Adobe dominate with multi-million-dollar budgets. Compete on specificity: regulated industry, authentication, compliance, and niche competitor displacement.")

add_h3(doc, "Bucket 1 — Competitor displacement (~$800–1,000/month)")
add_body(doc, "Target: FI evaluators actively comparing solutions")
add_table(doc,
    headers=["Keyword", "Intent level"],
    rows=[
        ["docusign alternative for banks",               "High"],
        ["docusign alternative financial services",      "High"],
        ["docusign alternative credit unions",           "High — lower CPC than generic"],
        ["adobe sign alternative compliance",            "High"],
        ["docusign competitors banking",                 "High"],
        ["better esign for financial institutions",      "Medium"],
    ],
    col_widths=[3.5, 3.0]
)
add_body(doc, "Match type: Phrase and exact. Broad match excluded — attracts unqualified traffic on DocuSign brand terms.")

add_h3(doc, "Bucket 2 — Industry / use-case specific (~$900–1,100/month)")
add_body(doc, "Target: FI ops, lending, and IT evaluators with active workflow problems")
add_table(doc,
    headers=["Keyword", "Intent level"],
    rows=[
        ["electronic signature for banks",               "High"],
        ["e-signature for credit unions",                "High — lower competition"],
        ["loan document electronic signature",           "High — very specific"],
        ["electronic signature lending compliance",      "High"],
        ["e-signature for financial institutions",       "High"],
        ["digital signature mortgage documents",         "Medium"],
        ["electronic signature NCUA compliant",          "High intent, very low competition"],
        ["electronic signature for community banks",     "Medium — niche, lower CPC"],
    ],
    col_widths=[3.5, 3.0]
)

add_h3(doc, "Bucket 3 — Authentication and security (~$500–700/month)")
add_body(doc, "Target: FI ops and compliance leaders thinking about fraud risk; dovetails with Mike's email campaign")
add_table(doc,
    headers=["Keyword", "Intent level"],
    rows=[
        ["e-signature with identity verification",           "High — core SIGNiX differentiator"],
        ["electronic signature signer authentication",       "High"],
        ["KBA electronic signature",                         "High intent, very low competition"],
        ["electronic signature fraud prevention banking",    "High"],
        ["two-factor authentication electronic signature",   "High"],
        ["biometric e-signature financial services",         "Medium — emerging search behavior"],
    ],
    col_widths=[3.5, 3.0]
)

add_h3(doc, "Bucket 4 — API and developer terms (~$400–600/month)")
add_body(doc, "Target: ISV product managers and developers evaluating Flex API; complements LinkedIn ISV targeting")
add_table(doc,
    headers=["Keyword", "Intent level"],
    rows=[
        ["e-signature API financial services",       "High"],
        ["document signing API for banks",           "High"],
        ["white label e-signature API",              "High — very low competition"],
        ["electronic signature SDK compliance",      "High intent, low volume"],
        ["sign document API fintech",                "High"],
        ["embedded e-signature API",                 "High"],
    ],
    col_widths=[3.5, 3.0]
)

add_h3(doc, "Bucket 5 — RON institutional only (~$200–300/month)")
add_body(doc, "Drop all individual notary terms. Keep institutional/platform-level terms only:")
add_table(doc,
    headers=["Keyword", "Intent level"],
    rows=[
        ["remote online notary platform financial institutions", "High — institutional buyers only"],
        ["RON platform compliance",                              "High"],
    ],
    col_widths=[3.5, 3.0]
)

add_h3(doc, "Ad message principles")
add_bullet(doc, "Lead with PKI, authentication, and regulated industry specificity — not 'easy' or 'fast' signing")
add_bullet(doc, "Competitor displacement ads: lead with compliance specificity or pricing transparency")
add_bullet(doc, "Each bucket should land on a dedicated landing page tailored to that intent — not the generic homepage")
add_bullet(doc, "Authentication bucket: connect to fraud risk; tie to Mike's email campaign messaging")
add_hr(doc)

# ── Section 4 ─────────────────────────────────────────────────────────────────
add_h2(doc, "4. LinkedIn Strategy ($1,500/month)")
add_h3(doc, "Why LinkedIn for this audience")
add_body(doc, "FI decision-makers and ISV product leaders do not frequently Google for e-sign solutions unprompted. They respond to peer content, authority positioning, and problem-aware messaging when it appears in their professional feed. LinkedIn is the right channel to build awareness and trust before a search event occurs.")

add_h3(doc, "Audience 1 — FI operations and compliance ($800/month)")
add_table(doc,
    headers=["Targeting dimension", "Values"],
    rows=[
        ["Job titles",    "VP Operations, Director of Digital Banking, IT Director, Lending Operations Manager, Chief Compliance Officer, VP Retail Banking, Director of Lending"],
        ["Industries",    "Banking, Financial Services, Credit Unions"],
        ["Company size",  "50–5,000 employees"],
        ["Geography",     "United States"],
        ["Message angle", "Show them how SIGNiX stops fraud and proves compliance"],
    ],
    col_widths=[2.0, 4.5]
)

add_h3(doc, "Audience 2 — ISV and Flex API targets ($700/month)")
add_table(doc,
    headers=["Targeting dimension", "Values"],
    rows=[
        ["Job titles",    "Product Manager, VP of Product, CTO, VP Engineering, Head of Product, Director of Integrations"],
        ["Industries",    "Financial Services Software, Computer Software, Fintech, Information Technology"],
        ["Company size",  "10–500 employees"],
        ["Geography",     "United States"],
        ["Message angle", "Show them a white-label API built for regulated industries"],
    ],
    col_widths=[2.0, 4.5]
)

add_h3(doc, "Format and approach")
add_bullet(doc, "Sponsored content (single image or document ads) — not lead gen forms in months 1–2")
add_bullet(doc, "Drive to dedicated landing pages, not the homepage")
add_bullet(doc, "Goal in months 1–2: audience validation — which persona engages more, then optimize budget toward the higher performer")
add_bullet(doc, "LinkedIn CPCs for B2B: $8–20+ per click; volume will be modest at $1,500/month; quality over quantity")
add_hr(doc)

# ── Section 5 ─────────────────────────────────────────────────────────────────
add_h2(doc, "5. Landing Pages (prerequisite for launch)")
add_body(doc, "Each keyword bucket and LinkedIn audience should land on a page that matches the search intent or ad message. Sending all paid traffic to the homepage is the single biggest waste of paid media budget.")
add_table(doc,
    headers=["Traffic source", "Recommended landing page focus"],
    rows=[
        ["Competitor displacement", "How SIGNiX compares to DocuSign for financial institutions"],
        ["Industry / use-case",     "Electronic signature for banks and credit unions — compliance-ready"],
        ["Authentication bucket",   "Signer authentication built in — not bolted on"],
        ["API / developer",         "SIGNiX Flex API — embed e-sign in your financial platform"],
        ["LinkedIn FI audience",    "Same as authentication or industry landing page"],
        ["LinkedIn ISV audience",   "Same as Flex API landing page"],
    ],
    col_widths=[2.5, 4.0]
)
add_body(doc, "If dedicated pages don't exist yet, prioritize building two: one for FI buyers (authentication/compliance angle) and one for ISV developers (Flex API).")
add_hr(doc)

# ── Section 6 ─────────────────────────────────────────────────────────────────
add_h2(doc, "6. Leading and Lagging Indicators")
add_h3(doc, "Leading indicators — check weekly")
add_table(doc,
    headers=["Indicator", "What it tells you", "Target"],
    rows=[
        ["Impression share by keyword bucket",          "Are we visible on priority terms?",           ">40% on priority buckets"],
        ["CTR by bucket",                               "Are the right people clicking?",              "3–6% for B2B search"],
        ["Average CPC trends",                          "Is spend efficient or inflating?",            "Monitor week-over-week"],
        ["Landing page bounce rate",                    "Is traffic finding what they expected?",      "<60%"],
        ["Form fill / demo request rate",               "Are visitors converting?",                    "2–5% of paid clicks"],
        ["LinkedIn engagement rate",                    "Are the right people stopping?",              ">0.5% sponsored content"],
        ["Lead qualification rate",                     "Are leads matching our target buyer?",        ">40% qualified"],
        ["Time from lead to first sales touchpoint",    "Is the handoff working?",                     "<48 hours"],
    ],
    col_widths=[2.2, 2.3, 2.0]
)

add_h3(doc, "Lagging indicators — check monthly and quarterly")
add_table(doc,
    headers=["Indicator", "What it tells you", "Notes"],
    rows=[
        ["Cost per qualified lead (CPL) by channel",       "Which channel delivers the most efficient leads?",  "Split by Google bucket + LinkedIn audience"],
        ["MQLs per month",                                 "Is pipeline building?",                             "Define MQL criteria before launch"],
        ["Sales-accepted leads (SALs)",                    "Is the sales team seeing quality?",                 "Track rep feedback per lead source"],
        ["Pipeline sourced ($) from paid channels",        "What's the revenue potential?",                     "Attribute in HubSpot"],
        ["Deals closed from paid sources",                 "What's the actual return?",                         "90–180 day lag expected"],
        ["Average deal ACV by lead source",                "Are paid leads higher value?",                      "Compare vs. notary lead ACV"],
        ["Staff hours per closed deal by source",          "The number that shows real cost",                   "Estimate with rep + CSM time logs"],
        ["Transaction volume from new paid-sourced customers", "Does it generate recurring revenue?",           "The ultimate SIGNiX metric"],
    ],
    col_widths=[2.2, 2.3, 2.0]
)
add_hr(doc)

# ── Section 7 ─────────────────────────────────────────────────────────────────
add_h2(doc, "7. 90-Day Launch Plan")
add_table(doc,
    headers=["Week", "Action", "Owner"],
    rows=[
        ["1–2",  "Pause all individual notary keyword spending; maintain RON institutional terms",              "Chris"],
        ["1–2",  "Build or identify landing pages for FI buyer + Flex API audiences",                          "Chris + web"],
        ["3–4",  "Launch Google Ads Buckets 1–3 at 80% of allocation; hold back 20% for optimization",        "Chris"],
        ["3–4",  "Launch LinkedIn Audience 1 (FI ops/compliance) with 2 ad variants",                         "Chris"],
        ["5–6",  "Review CTR and qualification rate; pause underperforming ad variants",                       "Chris"],
        ["6–8",  "Launch Google Bucket 4 (API/developer); launch LinkedIn Audience 2 (ISV)",                  "Chris"],
        ["8–12", "Review CPL by channel; reallocate toward top performer",                                     "Chris"],
        ["12",   "First 90-day report: CPL, pipeline sourced, SALs, ACV vs. notary benchmark",                "Chris to CEO"],
    ],
    col_widths=[0.6, 5.2, 0.8]
)
add_hr(doc)

# ── Section 8 ─────────────────────────────────────────────────────────────────
add_h2(doc, "8. Rationale Summary")
add_blockquote(doc,
    "We are currently spending $5,000/month generating leads that require too much staff time to close and onboard, "
    "and produce low recurring transaction volume. We are proposing to redirect that same budget toward FI operations "
    "and IT buyers who are actively evaluating e-sign and authentication solutions, and toward ISV developers who can "
    "embed SIGNiX at scale via the Flex API. The shift is revenue-neutral in spend, but the average deal value and "
    "recurring transaction potential per new customer is significantly higher — and we will be able to measure it "
    "against cost per qualified lead within 60 days."
)
add_hr(doc)
add_body(doc, "Last updated: April 2026. Update as campaigns launch and performance data is available.", italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/chris/Desktop/AI Summit 2026 Q2/proj-template-and-lease-SIGNiX-app/PROJECT-DOCS/SIGNIX-PAID-MEDIA-PLAN-2026.docx"
doc.save(out)
print(f"Saved: {out}")
