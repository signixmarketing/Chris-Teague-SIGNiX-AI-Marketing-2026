#!/usr/bin/env python3
"""Build SIGNiX_ABM_Scorecard.docx — weekly ABM activity tracker for sales team."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colours ─────────────────────────────────────────────────────────────
GREEN = RGBColor(0x6d, 0xa3, 0x4a)
INK   = RGBColor(0x2e, 0x34, 0x40)
BODY  = RGBColor(0x54, 0x54, 0x54)
MUTED = RGBColor(0x6b, 0x72, 0x80)
WHITE = RGBColor(0xff, 0xff, 0xff)
FONT  = "Calibri"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def style_run(run, color, bold=False, size=10, italic=False):
    run.font.name      = FONT
    run.font.color.rgb = color
    run.font.bold      = bold
    run.font.size      = Pt(size)
    run.font.italic    = italic

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=18)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, INK, bold=True, size=12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "6da34a")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_body(doc, text, size=9, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, BODY, size=size, italic=italic)
    p.paragraph_format.space_after = Pt(3)

def add_bullet(doc, text, size=9):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    style_run(run, BODY, size=size)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "d8dee9")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ── BUILD ─────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(11)      # Landscape
sec.page_height   = Inches(8.5)
sec.left_margin   = Inches(0.5)
sec.right_margin  = Inches(0.5)
sec.top_margin    = Inches(0.6)
sec.bottom_margin = Inches(0.6)

# Set landscape orientation
pgSz = OxmlElement('w:pgSz')
pgSz.set(qn('w:w'), '15840')   # 11 inches in twips
pgSz.set(qn('w:h'), '12240')   # 8.5 inches in twips
pgSz.set(qn('w:orient'), 'landscape')
sec._sectPr.append(pgSz)

doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(9)

# ── Title block ───────────────────────────────────────────────────────────────
add_h1(doc, "SIGNiX — ABM Activity Scorecard")

meta_rows = [
    ("Week of",      "____________"),
    ("Reviewed by",  "Chris + Sales Leader"),
    ("Review cadence", "Weekly — Monday morning"),
    ("Purpose",      "Make ABM outreach visible before pipeline shows up. Track accounts, activity, and status in one place."),
]
mt = doc.add_table(rows=1, cols=4)
mt.style = "Table Grid"
for i, (k, v) in enumerate(meta_rows):
    if i >= len(mt.rows[0].cells):
        break
    # Add cols as needed — build as 4-col row
for i, (k, v) in enumerate(meta_rows):
    # Use a simple 2-col layout repeated side by side
    pass

# Simple meta bar
meta_t = doc.add_table(rows=2, cols=4)
meta_t.style = "Table Grid"
meta_pairs = [("Week of", "____________"), ("Reviewed by", "Chris + Sales Leader"),
              ("Review cadence", "Weekly — Monday"), ("Purpose", "Make outreach visible before pipeline arrives")]
for i, (k, v) in enumerate(meta_pairs):
    set_cell_bg(meta_t.rows[0].cells[i], "2e3440")
    r_k = meta_t.rows[0].cells[i].paragraphs[0].add_run(k)
    style_run(r_k, WHITE, bold=True, size=8)
    set_cell_bg(meta_t.rows[1].cells[i], "f8fafb")
    r_v = meta_t.rows[1].cells[i].paragraphs[0].add_run(v)
    style_run(r_v, BODY, size=8)

for row in meta_t.rows:
    for i, w in enumerate([Inches(0.9), Inches(1.4), Inches(1.2), Inches(3.5)]):
        row.cells[i].width = w

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_hr(doc)

# ── Main scorecard table ──────────────────────────────────────────────────────
add_h2(doc, "Account Activity Tracker")

headers = [
    "Account Name",
    "Segment",
    "Owner",
    "First Contact",
    "Touches This Week\n(Call / Email / LI)",
    "Meeting\nBooked?",
    "Meeting\nDate",
    "Content Sent",
    "Status",
    "Notes / Next Step",
]

# 35 blank data rows
blank_row = [""] * len(headers)
rows = [blank_row[:] for _ in range(35)]

tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
tbl.style = "Table Grid"

# Header row
for ci, h in enumerate(headers):
    c = tbl.rows[0].cells[ci]
    set_cell_bg(c, "2e3440")
    c.paragraphs[0].clear()
    run = c.paragraphs[0].add_run(h)
    style_run(run, WHITE, bold=True, size=8)

# Data rows — alternating background
for r_i, row in enumerate(rows):
    bg_hex = "f8fafb" if r_i % 2 == 0 else "ffffff"
    for c_i, val in enumerate(row):
        c = tbl.rows[r_i + 1].cells[c_i]
        set_cell_bg(c, bg_hex)
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(val)
        style_run(run, BODY, size=8)
        c.paragraphs[0].paragraph_format.space_after = Pt(1)

# Column widths — landscape 10" usable
col_widths = [
    Inches(1.5),   # Account Name
    Inches(0.85),  # Segment
    Inches(0.7),   # Owner
    Inches(0.7),   # First Contact
    Inches(1.0),   # Touches
    Inches(0.6),   # Meeting booked
    Inches(0.6),   # Meeting date
    Inches(1.0),   # Content sent
    Inches(0.75),  # Status
    Inches(2.3),   # Notes
]
for row in tbl.rows:
    for ci, cw in enumerate(col_widths):
        row.cells[ci].width = cw

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_hr(doc)

# ── Reference section ─────────────────────────────────────────────────────────
add_h2(doc, "Reference — Status Definitions and Usage Guide")

# Two-column reference layout
ref_t = doc.add_table(rows=1, cols=2)
ref_t.style = "Table Grid"

# Left: Status definitions
lc = ref_t.rows[0].cells[0]
set_cell_bg(lc, "f8fafb")
lc.width = Inches(3.5)
lp = lc.paragraphs[0]
r = lp.add_run("STATUS DEFINITIONS")
style_run(r, GREEN, bold=True, size=8)

statuses = [
    ("Cold",               "No response yet. Account is on the list but not yet engaged."),
    ("Engaged",            "Account has responded, opened content, or shown any signal of interest."),
    ("Meeting Scheduled",  "A discovery or demo call is booked with a contact at this account."),
    ("Proposal",           "A formal proposal or pricing conversation has started."),
    ("Closed — Won",       "Deal signed. Move to customer success handoff."),
    ("Closed — Lost",      "No opportunity at this time. Note reason. Revisit in 90 days."),
    ("On Hold",            "Decision delayed. Follow up date noted in Notes column."),
]
for status, defn in statuses:
    p2 = lc.add_paragraph()
    r_s = p2.add_run(status + ":  ")
    style_run(r_s, INK, bold=True, size=8)
    r_d = p2.add_run(defn)
    style_run(r_d, BODY, size=8)
    p2.paragraph_format.space_after = Pt(2)

# Right: Usage guide
rc = ref_t.rows[0].cells[1]
set_cell_bg(rc, "f8fafb")
rc.width = Inches(6.5)
rp = rc.paragraphs[0]
r2 = rp.add_run("HOW TO USE THIS SCORECARD")
style_run(r2, GREEN, bold=True, size=8)

usage = [
    ("Update weekly",        "Every rep updates their accounts by end of day Friday. Review together Monday morning."),
    ("Touches column",       "Log the number of calls, emails, and LinkedIn messages sent that week. Format: 2C / 1E / 1LI."),
    ("Content sent",         "Note what you sent: Fraud Alert brief, HIPAA one-pager, competitor comparison, case study, etc."),
    ("Meeting booked",       "Mark Y when a meeting is confirmed. Add the date in the Meeting Date column."),
    ("Notes / Next step",    "One sentence only. What happens next and when. Example: 'Follow up Dec 15 after budget review.'"),
    ("Segment values",       "Use: Fin Services / Healthcare / Wealth Mgmt / Legal-Collections / Other"),
    ("Target account count", "Keep 30–40 active accounts on this list at all times. Add new accounts as others close or go inactive."),
]
for label, guide in usage:
    p3 = rc.add_paragraph()
    r_l = p3.add_run(label + ":  ")
    style_run(r_l, INK, bold=True, size=8)
    r_g = p3.add_run(guide)
    style_run(r_g, BODY, size=8)
    p3.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body(doc, "SIGNiX  ·  ABM Activity Scorecard  ·  April 2026  ·  Update weekly, review Monday", size=8, italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/chris/Desktop/AI Summit 2026 Q2/proj-template-and-lease-SIGNiX-app/PROJECT-DOCS/SIGNiX_ABM_Scorecard.docx"
doc.save(out)
print(f"Saved: {out}")
